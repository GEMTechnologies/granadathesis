#!/usr/bin/env python3
"""
Export Utilities - Multiple Format Support

Formats supported:
- BibTeX (for LaTeX)
- RIS (for Zotero - already implemented)
- EndNote XML
- CSV (for Excel)
- Excel (.xlsx)
- Word (.docx) bibliography
- Markdown tables
- JSON (already implemented)
"""

import csv
import json
from typing import List, Dict, Any
from pathlib import Path
from datetime import datetime

# Excel support
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Word support
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


class ExportService:
    """Export papers to multiple formats."""
    
    @staticmethod
    def to_bibtex(papers: List[Dict[str, Any]]) -> str:
        """
        Export to BibTeX format.
        
        Args:
            papers: List of paper dictionaries
            
        Returns:
            BibTeX formatted string
        """
        bibtex_entries = []
        
        for i, paper in enumerate(papers, 1):
            # Generate citation key
            authors = paper.get("authors", [])
            first_author = authors[0].split()[-1] if authors else "Unknown"
            year = paper.get("year", "")
            title_words = paper.get("title", "").split()[:2]
            cite_key = f"{first_author}{year}{''.join(title_words)}"
            cite_key = "".join(c for c in cite_key if c.isalnum())
            
            # Build entry
            entry = f"@article{{{cite_key},\n"
            entry += f"  title = {{{paper.get('title', '')}}},\n"
            
            # Authors
            if authors:
                author_str = " and ".join(authors)
                entry += f"  author = {{{author_str}}},\n"
            
            # Year
            if year:
                entry += f"  year = {{{year}}},\n"
            
            # Journal/Venue
            if paper.get("venue"):
                entry += f"  journal = {{{paper['venue']}}},\n"
            
            # DOI
            if paper.get("doi"):
                entry += f"  doi = {{{paper['doi']}}},\n"
            
            # URL
            if paper.get("url"):
                entry += f"  url = {{{paper['url']}}},\n"
            
            # Abstract
            if paper.get("abstract"):
                abstract = paper['abstract'].replace("{", "\\{").replace("}", "\\}")
                entry += f"  abstract = {{{abstract}}},\n"
            
            entry += "}\n"
            bibtex_entries.append(entry)
        
        return "\n".join(bibtex_entries)
    
    @staticmethod
    def to_endnote_xml(papers: List[Dict[str, Any]]) -> str:
        """
        Export to EndNote XML format.
        
        Args:
            papers: List of paper dictionaries
            
        Returns:
            EndNote XML formatted string
        """
        xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
        xml += '<xml>\n<records>\n'
        
        for paper in papers:
            xml += '  <record>\n'
            xml += '    <database name="Academic Search" path="AutoGranada">Academic Search.enl</database>\n'
            xml += '    <source-app name="AutoGranada">AutoGranada</source-app>\n'
            xml += '    <rec-number>1</rec-number>\n'
            xml += '    <ref-type name="Journal Article">17</ref-type>\n'
            
            # Title
            xml += f'    <titles>\n'
            xml += f'      <title>{paper.get("title", "")}</title>\n'
            xml += f'    </titles>\n'
            
            # Authors
            if paper.get("authors"):
                xml += '    <contributors>\n'
                xml += '      <authors>\n'
                for author in paper["authors"]:
                    xml += f'        <author>{author}</author>\n'
                xml += '      </authors>\n'
                xml += '    </contributors>\n'
            
            # Year
            if paper.get("year"):
                xml += f'    <dates>\n'
                xml += f'      <year>{paper["year"]}</year>\n'
                xml += f'    </dates>\n'
            
            # Journal
            if paper.get("venue"):
                xml += f'    <periodical>\n'
                xml += f'      <full-title>{paper["venue"]}</full-title>\n'
                xml += f'    </periodical>\n'
            
            # DOI
            if paper.get("doi"):
                xml += f'    <electronic-resource-num>{paper["doi"]}</electronic-resource-num>\n'
            
            # URL
            if paper.get("url"):
                xml += f'    <urls>\n'
                xml += f'      <related-urls>\n'
                xml += f'        <url>{paper["url"]}</url>\n'
                xml += f'      </related-urls>\n'
                xml += f'    </urls>\n'
            
            # Abstract
            if paper.get("abstract"):
                xml += f'    <abstract>{paper["abstract"]}</abstract>\n'
            
            xml += '  </record>\n'
        
        xml += '</records>\n</xml>'
        return xml
    
    @staticmethod
    def to_csv(papers: List[Dict[str, Any]], output_path: Path):
        """
        Export to CSV format.
        
        Args:
            papers: List of paper dictionaries
            output_path: Path to save CSV file
        """
        if not papers:
            return
        
        # Define columns
        columns = ["Title", "Authors", "Year", "Venue", "Citations", "DOI", "URL", "Source", "Abstract"]
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=columns)
            writer.writeheader()
            
            for paper in papers:
                writer.writerow({
                    "Title": paper.get("title", ""),
                    "Authors": "; ".join(paper.get("authors", [])),
                    "Year": paper.get("year", ""),
                    "Venue": paper.get("venue", ""),
                    "Citations": paper.get("citations", 0),
                    "DOI": paper.get("doi", ""),
                    "URL": paper.get("url", ""),
                    "Source": paper.get("source", ""),
                    "Abstract": paper.get("abstract", "")[:500]  # Truncate long abstracts
                })
    
    @staticmethod
    def to_excel(papers: List[Dict[str, Any]], output_path: Path):
        """
        Export to Excel format with formatting.
        
        Args:
            papers: List of paper dictionaries
            output_path: Path to save Excel file
        """
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl not installed - cannot create Excel files")
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Papers"
        
        # Headers
        headers = ["#", "Title", "Authors", "Year", "Venue", "Citations", "DOI", "URL", "Source"]
        ws.append(headers)
        
        # Format headers
        for cell in ws[1]:
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Data rows
        for i, paper in enumerate(papers, 1):
            ws.append([
                i,
                paper.get("title", ""),
                ", ".join(paper.get("authors", [])[:3]),  # First 3 authors
                paper.get("year", ""),
                paper.get("venue", ""),
                paper.get("citations", 0),
                paper.get("doi", ""),
                paper.get("url", ""),
                paper.get("source", "")
            ])
        
        # Adjust column widths
        ws.column_dimensions['B'].width = 50  # Title
        ws.column_dimensions['C'].width = 30  # Authors
        ws.column_dimensions['E'].width = 30  # Venue
        ws.column_dimensions['H'].width = 40  # URL
        
        wb.save(output_path)
    
    @staticmethod
    def to_word(papers: List[Dict[str, Any]], output_path: Path):
        """
        Export to Word document with formatted bibliography.
        
        Args:
            papers: List of paper dictionaries
            output_path: Path to save Word file
        """
        if not WORD_AVAILABLE:
            raise ImportError("python-docx not installed - cannot create Word files")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Academic Paper Bibliography', 0)
        title.alignment = 1  # Center
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Total Papers: {len(papers)}")
        doc.add_paragraph("")
        
        # Papers
        for i, paper in enumerate(papers, 1):
            # Number
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. ")
            run.bold = True
            
            # Title
            run = p.add_run(paper.get("title", ""))
            run.bold = True
            run.font.size = Pt(12)
            
            # Authors
            authors = paper.get("authors", [])
            if authors:
                p = doc.add_paragraph()
                p.add_run("Authors: ").italic = True
                p.add_run(", ".join(authors))
            
            # Year and Venue
            details = []
            if paper.get("year"):
                details.append(f"Year: {paper['year']}")
            if paper.get("venue"):
                details.append(f"Venue: {paper['venue']}")
            if paper.get("citations"):
                details.append(f"Citations: {paper['citations']}")
            
            if details:
                p = doc.add_paragraph(" | ".join(details))
                p.style = 'List Bullet'
            
            # DOI and URL
            if paper.get("doi"):
                p = doc.add_paragraph()
                p.add_run("DOI: ").italic = True
                p.add_run(paper["doi"])
            
            if paper.get("url"):
                p = doc.add_paragraph()
                p.add_run("URL: ").italic = True
                run = p.add_run(paper["url"])
                run.font.color.rgb = RGBColor(0, 0, 255)
            
            # Abstract (truncated)
            if paper.get("abstract"):
                p = doc.add_paragraph()
                p.add_run("Abstract: ").italic = True
                abstract = paper["abstract"][:300] + "..." if len(paper["abstract"]) > 300 else paper["abstract"]
                p.add_run(abstract)
            
            # Separator
            doc.add_paragraph("_" * 80)
        
        doc.save(output_path)
    
    @staticmethod
    def to_markdown(papers: List[Dict[str, Any]]) -> str:
        """
        Export to Markdown table format.
        
        Args:
            papers: List of paper dictionaries
            
        Returns:
            Markdown formatted string
        """
        md = "# Academic Papers\n\n"
        md += f"**Total:** {len(papers)} papers\n\n"
        
        # Table
        md += "| # | Title | Authors | Year | Citations | Source |\n"
        md += "|---|-------|---------|------|-----------|--------|\n"
        
        for i, paper in enumerate(papers, 1):
            title = paper.get("title", "")[:50] + "..." if len(paper.get("title", "")) > 50 else paper.get("title", "")
            authors = ", ".join(paper.get("authors", [])[:2])
            year = paper.get("year", "")
            citations = paper.get("citations", 0)
            source = paper.get("source", "")
            
            md += f"| {i} | {title} | {authors} | {year} | {citations} | {source} |\n"
        
        md += "\n## Detailed List\n\n"
        
        for i, paper in enumerate(papers, 1):
            md += f"### {i}. {paper.get('title', '')}\n\n"
            md += f"**Authors:** {', '.join(paper.get('authors', []))}\n\n"
            md += f"**Year:** {paper.get('year', 'N/A')} | **Citations:** {paper.get('citations', 0)} | **Source:** {paper.get('source', '')}\n\n"
            
            if paper.get("venue"):
                md += f"**Venue:** {paper['venue']}\n\n"
            
            if paper.get("doi"):
                md += f"**DOI:** {paper['doi']}\n\n"
            
            if paper.get("url"):
                md += f"**URL:** [{paper['url']}]({paper['url']})\n\n"
            
            if paper.get("abstract"):
                md += f"**Abstract:** {paper['abstract'][:200]}...\n\n"
            
            md += "---\n\n"
        
        return md


def export_papers(papers: List[Dict[str, Any]], output_dir: Path, base_name: str):
    """
    Export papers to all available formats.
    
    Args:
        papers: List of paper dictionaries
        output_dir: Output directory
        base_name: Base filename (without extension)
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    exporter = ExportService()
    results = {}
    
    # BibTeX
    try:
        bibtex_path = output_dir / f"{base_name}.bib"
        bibtex_path.write_text(exporter.to_bibtex(papers))
        results["bibtex"] = str(bibtex_path)
        print(f"✓ Exported BibTeX: {bibtex_path.name}")
    except Exception as e:
        print(f"✗ BibTeX export failed: {e}")
    
    # EndNote XML
    try:
        endnote_path = output_dir / f"{base_name}.xml"
        endnote_path.write_text(exporter.to_endnote_xml(papers))
        results["endnote"] = str(endnote_path)
        print(f"✓ Exported EndNote XML: {endnote_path.name}")
    except Exception as e:
        print(f"✗ EndNote export failed: {e}")
    
    # CSV
    try:
        csv_path = output_dir / f"{base_name}.csv"
        exporter.to_csv(papers, csv_path)
        results["csv"] = str(csv_path)
        print(f"✓ Exported CSV: {csv_path.name}")
    except Exception as e:
        print(f"✗ CSV export failed: {e}")
    
    # Excel
    if EXCEL_AVAILABLE:
        try:
            excel_path = output_dir / f"{base_name}.xlsx"
            exporter.to_excel(papers, excel_path)
            results["excel"] = str(excel_path)
            print(f"✓ Exported Excel: {excel_path.name}")
        except Exception as e:
            print(f"✗ Excel export failed: {e}")
    
    # Word
    if WORD_AVAILABLE:
        try:
            word_path = output_dir / f"{base_name}.docx"
            exporter.to_word(papers, word_path)
            results["word"] = str(word_path)
            print(f"✓ Exported Word: {word_path.name}")
        except Exception as e:
            print(f"✗ Word export failed: {e}")
    
    # Markdown
    try:
        md_path = output_dir / f"{base_name}.md"
        md_path.write_text(exporter.to_markdown(papers))
        results["markdown"] = str(md_path)
        print(f"✓ Exported Markdown: {md_path.name}")
    except Exception as e:
        print(f"✗ Markdown export failed: {e}")
    
    # JSON
    try:
        json_path = output_dir / f"{base_name}.json"
        json_path.write_text(json.dumps(papers, indent=2))
        results["json"] = str(json_path)
        print(f"✓ Exported JSON: {json_path.name}")
    except Exception as e:
        print(f"✗ JSON export failed: {e}")
    
    return results
