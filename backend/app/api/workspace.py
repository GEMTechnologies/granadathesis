"""
Workspace API Endpoints
Handles user workspaces, folders, files, and projects management
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Response
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from pathlib import Path
import os
import uuid
import shutil
import aiofiles
import re
import io
from docx import Document
from docx.shared import Inches
from datetime import datetime


router = APIRouter(prefix="/api/workspace", tags=["workspace"])

# Workspace base directory
# In Docker: use /app/workspaces (mounted from host)
# On host: use relative path from backend directory
if os.path.exists("/app/workspaces"):
    # Running in Docker
    WORKSPACES_BASE_DIR = Path("/app/workspaces")
elif os.path.exists("/home/gemtech/Desktop/thesis/workspaces"):
    # Running on host
    WORKSPACES_BASE_DIR = Path("/home/gemtech/Desktop/thesis/workspaces")
else:
    # Fallback: relative to this file
    WORKSPACES_BASE_DIR = Path(__file__).resolve().parent.parent.parent.parent / "workspaces"


# Request Models
class CreateFolderRequest(BaseModel):
    name: str
    parent_folder_id: Optional[str] = None


class CreateFileRequest(BaseModel):
    name: str
    folder_id: Optional[str] = None
    content: Optional[str] = ""


class CreateProjectRequest(BaseModel):
    name: str
    project_type: str = "thesis"
    folder_id: Optional[str] = None
    description: Optional[str] = None


class MoveItemRequest(BaseModel):
    target_folder_id: Optional[str] = None


class RenameItemRequest(BaseModel):
    item_id: str
    item_type: str
    new_name: str


# Helper Functions
def get_workspace_path(workspace_id: str) -> Path:
    """Get filesystem path for a workspace"""
    return WORKSPACES_BASE_DIR / workspace_id


# Default folders that are always created
DEFAULT_FOLDERS = [
    "sources",      # Research sources and papers
    "sections",     # Document sections
    "chapters",     # Thesis/document chapters
    "uploads",      # Uploaded files
    "outputs",      # Generated outputs
    "data",         # Data files
    "notes",        # Research notes
    "drafts",       # Draft documents
]

def ensure_workspace_exists(workspace_id: str) -> Path:
    """Ensure workspace directory exists with default folder structure"""
    try:
        # Ensure base directory exists first
        WORKSPACES_BASE_DIR.mkdir(parents=True, exist_ok=True)
        
        workspace_path = get_workspace_path(workspace_id)
        workspace_path.mkdir(parents=True, exist_ok=True)
        
        # Create default folders if they don't exist
        for folder_name in DEFAULT_FOLDERS:
            folder_path = workspace_path / folder_name
            folder_path.mkdir(parents=True, exist_ok=True)
            
            # Mark as system folder with metadata file
            metadata_file = folder_path / ".folder_info.json"
            if not metadata_file.exists():
                import json
                metadata = {
                    "name": folder_name,
                    "type": "system",
                    "default": True,
                    "created_at": datetime.now().isoformat(),
                    "description": _get_folder_description(folder_name)
                }
                metadata_file.write_text(json.dumps(metadata, indent=2))
        
        return workspace_path
    except Exception as e:
        error_msg = f"Failed to create workspace directory: {str(e)}. WORKSPACES_BASE_DIR: {WORKSPACES_BASE_DIR}"
        print(f"ERROR: {error_msg}")
        raise Exception(error_msg)

def _get_folder_description(folder_name: str) -> str:
    """Get description for default folders"""
    descriptions = {
        "sources": "Research sources, papers, and references",
        "sections": "Document sections and subsections",
        "chapters": "Thesis or document chapters",
        "uploads": "Uploaded files and documents",
        "outputs": "Generated outputs and exports",
        "data": "Data files and datasets",
        "notes": "Research notes and annotations",
        "drafts": "Draft documents and works in progress",
    }
    return descriptions.get(folder_name, "Workspace folder")


# Workspace Endpoints
@router.get("/{workspace_id}/structure")
async def get_workspace_structure(workspace_id: str):
    """Get complete workspace structure (folders, files, projects) - recursively scans all files"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        items = []
        
        def scan_directory(dir_path: Path, parent_path: str = ""):
            """Recursively scan directory and add all items"""
            if not dir_path.exists():
                return
            
            for item in dir_path.iterdir():
                if item.name.startswith('.'):
                    continue
                
                # Calculate relative path from workspace root
                relative_path = item.relative_to(workspace_path)
                item_path_str = str(relative_path)
                
                # Generate unique ID based on path
                item_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{workspace_id}/{item_path_str}"))
                
                # Determine parent path for nested items
                parent_item_path = str(relative_path.parent) if relative_path.parent != Path('.') else ""
                parent_item_id = None
                if parent_item_path:
                    parent_item_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{workspace_id}/{parent_item_path}"))
                
                if item.is_dir():
                    # Check if it's a project (has .project.json)
                    project_metadata = item / ".project.json"
                    folder_info = item / ".folder_info.json"
                    
                    item_type = "project" if project_metadata.exists() else "folder"
                    is_default = False
                    folder_description = None
                    project_type = None
                    
                    # Load project metadata if it exists
                    if project_metadata.exists():
                        try:
                            import json
                            proj_meta = json.loads(project_metadata.read_text())
                            project_type = proj_meta.get("project_type", "thesis")
                        except:
                            pass
                    
                    # Check if it's a default/system folder
                    if folder_info.exists():
                        try:
                            import json
                            info = json.loads(folder_info.read_text())
                            is_default = info.get("default", False)
                            folder_description = info.get("description")
                        except:
                            pass
                    elif item.name in DEFAULT_FOLDERS:
                        is_default = True
                    
                    item_data = {
                        "id": item_id,
                        "name": item.name,
                        "type": item_type,
                        "path": item_path_str,
                        "parentId": parent_item_id,
                        "isDefault": is_default,
                        "description": folder_description,
                        "url": f"/workspace/{workspace_id}/{item_path_str}",
                        "shareable": True,
                        "createdAt": datetime.fromtimestamp(item.stat().st_ctime).isoformat(),
                        "updatedAt": datetime.fromtimestamp(item.stat().st_mtime).isoformat(),
                    }
                    
                    # Add project type if it's a project
                    if project_type:
                        item_data["projectType"] = project_type
                    
                    items.append(item_data)
                    
                    # Recursively scan subdirectory
                    scan_directory(item, item_path_str)
                    
                elif item.is_file():
                    items.append({
                        "id": item_id,
                        "name": item.name,
                        "type": "file",
                        "path": item_path_str,
                        "parentId": parent_item_id,
                        "fileType": item.suffix[1:] if item.suffix else "unknown",
                        "size": item.stat().st_size,
                        "createdAt": datetime.fromtimestamp(item.stat().st_ctime).isoformat(),
                        "updatedAt": datetime.fromtimestamp(item.stat().st_mtime).isoformat(),
                    })
        
        # Start recursive scan from workspace root
        if workspace_path.exists():
            scan_directory(workspace_path)
        
        return {"items": items, "workspace_id": workspace_id}
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_detail = f"Error fetching workspace structure: {str(e)}\n{traceback.format_exc()}"
        print(f"ERROR in get_workspace_structure: {error_detail}")
        raise HTTPException(status_code=500, detail=f"Error fetching workspace structure: {str(e)}")


@router.post("/{workspace_id}/folders")
async def create_folder(workspace_id: str, request: CreateFolderRequest):
    """Create a new folder in workspace"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Validate folder name
        if not request.name or not request.name.strip():
            raise HTTPException(status_code=400, detail="Folder name is required")
        
        # Sanitize folder name - remove invalid characters
        folder_name = request.name.strip()
        # Replace invalid filesystem characters
        invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
        for char in invalid_chars:
            folder_name = folder_name.replace(char, '_')
        # Remove leading/trailing dots and spaces
        folder_name = folder_name.strip('. ')
        
        if not folder_name:
            raise HTTPException(status_code=400, detail="Invalid folder name")
        
        # Determine folder path
        if request.parent_folder_id:
            # Try to find parent folder
            # First check if it's a direct name match in workspace
            parent_path = workspace_path / request.parent_folder_id
            if not parent_path.exists() or not parent_path.is_dir():
                # If parent_folder_id is just a name, check if it exists
                parent_path = workspace_path / request.parent_folder_id
                if not parent_path.exists():
                    # Create in workspace root if parent not found
                    folder_path = workspace_path / folder_name
                else:
                    folder_path = parent_path / folder_name
            else:
                folder_path = parent_path / folder_name
        else:
            folder_path = workspace_path / folder_name
        
        if folder_path.exists():
            raise HTTPException(status_code=400, detail=f"Folder '{folder_name}' already exists")
        
        # Create folder
        folder_path.mkdir(parents=True, exist_ok=True)
        
        # Generate folder URL
        folder_url = f"/workspace/{workspace_id}/{folder_path.name}"
        
        return {
            "id": str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{workspace_id}/{folder_name}")),
            "name": folder_name,
            "path": str(folder_path.relative_to(workspace_path)),
            "type": "folder",
            "url": folder_url,
            "shareable_url": f"/shared/workspace/{workspace_id}/{folder_path.name}",
            "created_at": datetime.now().isoformat()
        }
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_detail = f"Error creating folder: {str(e)}\n{traceback.format_exc()}"
        print(f"ERROR in create_folder: {error_detail}")
        raise HTTPException(status_code=500, detail=f"Error creating folder: {str(e)}")


@router.post("/{workspace_id}/files")
async def create_file(workspace_id: str, request: CreateFileRequest):
    """Create a new file in workspace"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Validate file name
        if not request.name or not request.name.strip():
            raise HTTPException(status_code=400, detail="File name is required")
        
        file_name = request.name.strip()
        
        if request.folder_id:
            # Create in folder (simplified)
            folder_path = workspace_path / request.folder_id
            if folder_path.exists() and folder_path.is_dir():
                file_path = folder_path / file_name
            else:
                file_path = workspace_path / file_name
        else:
            file_path = workspace_path / file_name
        
        if file_path.exists():
            raise HTTPException(status_code=400, detail="File already exists")
        
        # Create file with content
        file_path.write_text(request.content or "", encoding='utf-8')
        
        return {
            "id": str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{workspace_id}/{file_name}")),
            "name": file_name,
            "path": str(file_path.relative_to(workspace_path)),
            "type": "file",
            "size": file_path.stat().st_size,
            "created_at": datetime.now().isoformat()
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating file: {str(e)}")


@router.post("/{workspace_id}/upload")
async def upload_file(
    workspace_id: str,
    file: UploadFile = File(...),
    folder_id: Optional[str] = None
):
    """Upload a file to workspace (auto-creates uploads folder if needed)"""
    workspace_path = ensure_workspace_exists(workspace_id)
    
    # Get uploads folder or use specified folder
    if folder_id:
        upload_folder = workspace_path / folder_id
    else:
        upload_folder = workspace_path / "uploads"
    
    upload_folder.mkdir(parents=True, exist_ok=True)
    
    # Save uploaded file
    file_path = upload_folder / file.filename
    async with aiofiles.open(file_path, 'wb') as f:
        content = await file.read()
        await f.write(content)
    
    return {
        "id": str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{workspace_id}/{file.filename}")),
        "name": file.filename,
        "path": str(file_path.relative_to(workspace_path)),
        "type": "file",
        "size": file_path.stat().st_size,
        "uploaded_at": datetime.now().isoformat()
    }


@router.post("/{workspace_id}/projects")
async def create_project(workspace_id: str, request: CreateProjectRequest):
    """Create a new project in workspace"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Validate project name
        if not request.name or not request.name.strip():
            raise HTTPException(status_code=400, detail="Project name is required")
        
        project_name = request.name.strip().replace('/', '_').replace('\\', '_')
        
        if request.folder_id:
            parent_path = workspace_path / request.folder_id
            if parent_path.exists() and parent_path.is_dir():
                project_path = parent_path / project_name
            else:
                project_path = workspace_path / project_name
        else:
            project_path = workspace_path / project_name
        
        if project_path.exists():
            raise HTTPException(status_code=400, detail="Project already exists")
        
        project_path.mkdir(parents=True, exist_ok=True)
        
        # Create project metadata file
        metadata = {
            "project_type": request.project_type,
            "description": request.description,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
        }
        
        metadata_file = project_path / ".project.json"
        import json
        metadata_file.write_text(json.dumps(metadata, indent=2))
        
        # Create default folders inside project
        for folder_name in DEFAULT_FOLDERS:
            (project_path / folder_name).mkdir(exist_ok=True)
        
        return {
            "id": str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{workspace_id}/{project_name}")),
            "name": project_name,
            "type": "project",
            "project_type": request.project_type,
            "path": str(project_path.relative_to(workspace_path)),
            "created_at": datetime.now().isoformat()
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating project: {str(e)}")


@router.delete("/{workspace_id}/folders/{folder_id}")
async def delete_folder(workspace_id: str, folder_id: str):
    """Delete a folder"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Find folder by name (simplified - in production would use folder_id from DB)
        folder_path = workspace_path / folder_id
        
        if not folder_path.exists() or not folder_path.is_dir():
            raise HTTPException(status_code=404, detail="Folder not found")
        
        # Don't allow deleting default/system folders
        if folder_id in DEFAULT_FOLDERS:
            raise HTTPException(status_code=400, detail="Cannot delete system folders")
        
        # Delete folder recursively
        import shutil
        shutil.rmtree(folder_path)
        
        return {"message": f"Folder '{folder_id}' deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting folder: {str(e)}")


@router.delete("/{workspace_id}/files/{file_id}")
async def delete_file(workspace_id: str, file_id: str):
    """Delete a file"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Find file by name (simplified)
        file_path = workspace_path / file_id
        
        if not file_path.exists() or not file_path.is_file():
            raise HTTPException(status_code=404, detail="File not found")
        
        file_path.unlink()
        
        return {"message": f"File '{file_id}' deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting file: {str(e)}")


@router.delete("/{workspace_id}/projects/{project_id}")
async def delete_project(workspace_id: str, project_id: str):
    """Delete a project"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Find project by name
        project_path = workspace_path / project_id
        
        if not project_path.exists() or not project_path.is_dir():
            raise HTTPException(status_code=404, detail="Project not found")
        
        # Delete project folder recursively
        import shutil
        shutil.rmtree(project_path)
        
        return {"message": f"Project '{project_id}' deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting project: {str(e)}")


@router.post("/{workspace_id}/folders/{folder_id}/move")
async def move_folder(workspace_id: str, folder_id: str, request: MoveItemRequest):
    """Move a folder to another location"""
    return {"message": "Move operation would be implemented with proper path resolution"}


@router.post("/{workspace_id}/files/{file_id}/move")
async def move_file(workspace_id: str, file_id: str, request: MoveItemRequest):
    """Move a file to another folder"""
    return {"message": "Move operation would be implemented with proper path resolution"}


@router.post("/{workspace_id}/rename")
async def rename_item(workspace_id: str, request: RenameItemRequest):
    """Rename a folder, file, or project"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Find item by name (simplified - uses item_id as name for now)
        old_path = workspace_path / request.item_id
        new_name = request.new_name.strip().replace('/', '_').replace('\\', '_')
        new_path = workspace_path / new_name
        
        if not old_path.exists():
            raise HTTPException(status_code=404, detail="Item not found")
        
        # Don't allow renaming default folders
        if request.item_id in DEFAULT_FOLDERS:
            raise HTTPException(status_code=400, detail="Cannot rename system folders")
        
        if new_path.exists():
            raise HTTPException(status_code=400, detail="A folder or file with this name already exists")
        
        # Rename
        old_path.rename(new_path)
        
        # Update project metadata if it's a project
        if request.item_type == 'project' and new_path.is_dir():
            metadata_file = new_path / ".project.json"
            if metadata_file.exists():
                import json
                metadata = json.loads(metadata_file.read_text())
                metadata["updated_at"] = datetime.now().isoformat()
                metadata_file.write_text(json.dumps(metadata, indent=2))
        
        return {
            "message": f"Renamed '{request.item_id}' to '{new_name}'",
            "new_name": new_name
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error renaming item: {str(e)}")


class ConvertToDocxRequest(BaseModel):
    content: str
    filename: str


@router.get("/{workspace_id}/serve/{path:path}")
async def serve_file(workspace_id: str, path: str):
    """Serve a file from the workspace"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        file_path = (workspace_path / path).resolve()
        
        # Security check: ensure file is within workspace
        if not str(file_path).startswith(str(workspace_path.resolve())):
             raise HTTPException(status_code=403, detail="Access denied")
             
        if not file_path.exists() or not file_path.is_file():
            raise HTTPException(status_code=404, detail="File not found")
            
        return FileResponse(file_path)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error serving file: {str(e)}")


@router.post("/{workspace_id}/convert/docx")
async def convert_to_docx(workspace_id: str, request: ConvertToDocxRequest):
    """Convert markdown content to DOCX with image embedding"""
    try:
        workspace_path = ensure_workspace_exists(workspace_id)
        
        # Create a new Document
        doc = Document()
        doc.add_heading(request.filename, 0)
        
        # Simple Markdown parsing (this is a basic implementation)
        # It handles headers, simple paragraphs, and images
        lines = request.content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Handle Headers
            if line.startswith('#'):
                level = len(line.split(' ')[0])
                text = line[level:].strip()
                if level <= 9:
                    doc.add_heading(text, level=min(level, 9))
                else:
                     doc.add_paragraph(text, style='Body Text')
            
            # Handle Images: ![alt](path)
            elif line.startswith('![') and '](' in line and line.endswith(')'):
                try:
                    # Extract path
                    match = re.search(r'\!\[.*?\]\((.*?)\)', line)
                    if match:
                        img_path_str = match.group(1)
                        # Resolve path relative to workspace
                        img_full_path = (workspace_path / img_path_str).resolve()
                        
                        if img_full_path.exists() and img_full_path.is_file():
                             doc.add_picture(str(img_full_path), width=Inches(6))
                        else:
                             doc.add_paragraph(f"[Image not found: {img_path_str}]")
                except Exception as img_err:
                     print(f"Error adding image: {str(img_err)}")
                     doc.add_paragraph(f"[Error loading image: {line}]")
            
            # Regular Paragraph
            else:
                doc.add_paragraph(line)
        
        # Save to buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Return as downloadable file
        headers = {
            'Content-Disposition': f'attachment; filename="{request.filename}.docx"'
        }
        return Response(
            content=file_stream.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers=headers
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error converting to DOCX: {str(e)}")
