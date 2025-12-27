"""
Base Thesis Generator - Core functionality for all universities
"""

from typing import Dict, List, Optional, Any
from abc import ABC, abstractmethod
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from pathlib import Path
from datetime import datetime


class BaseThesisGenerator(ABC):
    """Abstract base class for university-specific thesis generators"""

    def __init__(self, university_type: str, workspace_id: str):
        self.university_type = university_type
        self.workspace_id = workspace_id
        self.doc = Document()
        self.metadata = {
            "university_type": university_type,
            "workspace_id": workspace_id,
            "created_at": datetime.now().isoformat(),
            "chapters": {},
        }

    def generate_complete_thesis(self, thesis_input: Dict[str, Any]) -> Optional[str]:
        """
        Generate complete thesis from input data
        
        Args:
            thesis_input: Dictionary containing:
                - title: Thesis title
                - author_name: Author name
                - supervisor: Supervisor name
                - topic: Thesis topic/objectives
                - chapters: Dict of chapter content {chapter_num: content}
                - abstract: Abstract text
                - appendices: Optional appendices
        
        Returns:
            Path to generated .docx file or None
        """
        try:
            # Extract input data
            title = thesis_input.get("title", "")
            author = thesis_input.get("author_name", "")
            supervisor = thesis_input.get("supervisor", "")
            topic = thesis_input.get("topic", "")
            chapters = thesis_input.get("chapters", {})
            abstract = thesis_input.get("abstract", "")

            # Store metadata
            self.metadata.update({
                "title": title,
                "author": author,
                "supervisor": supervisor,
                "topic": topic,
            })

            # Generate thesis sections
            self._add_cover_page(title, author, supervisor)
            self._add_preliminary_pages(title, author, abstract)
            self._add_chapters(chapters)
            self._add_appendices(thesis_input.get("appendices", {}))

            # Save document
            output_path = self._save_thesis(f"thesis_{self.workspace_id}")
            
            # Save metadata
            self._save_metadata()

            return output_path

        except Exception as e:
            print(f"Error generating thesis: {str(e)}")
            return None

    @abstractmethod
    def _add_cover_page(self, title: str, author: str, supervisor: str):
        """Add cover page (university-specific implementation)"""
        pass

    def _add_preliminary_pages(self, title: str, author: str, abstract: str):
        """Add preliminary pages (can be overridden by subclasses)"""
        # Add approval page
        self._add_approval_page(author, title)
        
        # Add declaration page
        self._add_declaration_page(author)
        
        # Add dedication page
        self._add_dedication_page()
        
        # Add acknowledgement page
        self._add_acknowledgement_page()
        
        # Add abstract
        self._add_abstract(title, abstract)
        
        # Add table of contents
        self._add_table_of_contents()

    def _add_approval_page(self, author: str, title: str):
        """Add approval/certification page"""
        self.doc.add_paragraph()
        heading = self.doc.add_heading("Approval Page", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            f"This thesis titled \"{title}\" by {author} has been approved by the "
            "supervisory committee and is ready for submission."
        )
        self.doc.add_paragraph()
        self.doc.add_paragraph("_" * 50)
        self.doc.add_paragraph("Supervisor Name")
        self.doc.add_paragraph()
        self.doc.add_paragraph("_" * 50)
        self.doc.add_paragraph("Date")

    def _add_declaration_page(self, author: str):
        """Add declaration page"""
        self.doc.add_page_break()
        heading = self.doc.add_heading("Declaration", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        declaration_text = (
            f"I, {author}, hereby declare that this thesis is my own original work and has "
            "not been submitted in whole or in part to any other institution for the award of "
            "a degree. All sources cited have been acknowledged appropriately."
        )
        self.doc.add_paragraph(declaration_text)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph("_" * 50)
        self.doc.add_paragraph(f"{author}")
        self.doc.add_paragraph()
        self.doc.add_paragraph("_" * 50)
        self.doc.add_paragraph("Date")

    def _add_dedication_page(self):
        """Add dedication page"""
        self.doc.add_page_break()
        heading = self.doc.add_heading("Dedication", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        self.doc.add_paragraph("[Your dedication text here]")

    def _add_acknowledgement_page(self):
        """Add acknowledgement page"""
        self.doc.add_page_break()
        heading = self.doc.add_heading("Acknowledgements", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        self.doc.add_paragraph("[Your acknowledgements text here]")

    def _add_abstract(self, title: str, abstract: str):
        """Add abstract page"""
        self.doc.add_page_break()
        heading = self.doc.add_heading("Abstract", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        if abstract:
            self.doc.add_paragraph(abstract)
        else:
            self.doc.add_paragraph("[Your abstract text here - typically 250-500 words]")

    def _add_table_of_contents(self):
        """Add table of contents placeholder"""
        self.doc.add_page_break()
        heading = self.doc.add_heading("Table of Contents", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        self.doc.add_paragraph("[Table of Contents - Auto-generated]")

    def _add_chapters(self, chapters: Dict[int, str]):
        """Add main chapters"""
        for chapter_num in sorted(chapters.keys()):
            self.doc.add_page_break()
            
            chapter_content = chapters[chapter_num]
            if isinstance(chapter_content, str):
                # Simple text content
                heading = self.doc.add_heading(f"Chapter {chapter_num}", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.doc.add_paragraph()
                self.doc.add_paragraph(chapter_content)
                
                self.metadata["chapters"][chapter_num] = {
                    "title": f"Chapter {chapter_num}",
                    "length": len(chapter_content.split()),
                }
            elif isinstance(chapter_content, dict):
                # Structured chapter data
                title = chapter_content.get("title", f"Chapter {chapter_num}")
                content = chapter_content.get("content", "")
                
                heading = self.doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.doc.add_paragraph()
                self.doc.add_paragraph(content)
                
                self.metadata["chapters"][chapter_num] = {
                    "title": title,
                    "length": len(content.split()),
                }

    def _add_appendices(self, appendices: Dict[str, str]):
        """Add appendices"""
        if not appendices:
            return
            
        self.doc.add_page_break()
        heading = self.doc.add_heading("Appendices", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for app_title, app_content in appendices.items():
            self.doc.add_page_break()
            subheading = self.doc.add_heading(app_title, level=2)
            self.doc.add_paragraph()
            self.doc.add_paragraph(app_content)

    def _save_thesis(self, filename: str) -> str:
        """Save thesis document"""
        workspace_path = Path(f"/home/gemtech/Desktop/thesis/workspaces/{self.workspace_id}")
        workspace_path.mkdir(parents=True, exist_ok=True)
        
        file_path = workspace_path / f"{filename}.docx"
        self.doc.save(str(file_path))
        
        return str(file_path)

    def _save_metadata(self):
        """Save thesis metadata"""
        workspace_path = Path(f"/home/gemtech/Desktop/thesis/workspaces/{self.workspace_id}")
        metadata_path = workspace_path / "thesis_metadata.json"
        
        with open(metadata_path, "w") as f:
            json.dump(self.metadata, f, indent=2)

    def get_metadata(self) -> Dict[str, Any]:
        """Get thesis metadata"""
        return self.metadata
