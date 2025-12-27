"""
University of Juba PhD Thesis Generator

Specific implementation for University of Juba PhD students
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, Optional
from pathlib import Path
import json

from ..base_template.base_generator import BaseThesisGenerator


class UoJPhDGenerator(BaseThesisGenerator):
    """University of Juba PhD Thesis Generator"""

    def __init__(self, workspace_id: str):
        super().__init__("uoj_phd", workspace_id)
        self.university_name = "UNIVERSITY OF JUBA"
        self.school_name = "SCHOOL OF GRADUATE STUDIES"

    def _add_cover_page(self, title: str, author: str, supervisor: str):
        """Add University of Juba specific cover page"""
        # University name
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(self.university_name)
        run.font.size = Pt(12)
        run.font.bold = True

        # School name
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(self.school_name)
        run.font.size = Pt(12)
        run.font.bold = True

        # Department
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("DEPARTMENT OF [YOUR DEPARTMENT]")
        run.font.size = Pt(11)

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Title
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.font.size = Pt(14)
        run.font.bold = True

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Author
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"By\n{author}")
        run.font.size = Pt(12)

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Supervisor
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"Supervisor: {supervisor}")
        run.font.size = Pt(11)

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Date
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("[Month Year]")
        run.font.size = Pt(11)

    def _add_preliminary_pages(self, title: str, author: str, abstract: str):
        """Add UoJ-specific preliminary pages"""
        # Skip cover page (already added)
        
        # Add approval page
        self.doc.add_page_break()
        self._add_uoj_approval_page(author, title)
        
        # Add declaration page
        self.doc.add_page_break()
        self._add_uoj_declaration_page(author)
        
        # Add dedication page
        self.doc.add_page_break()
        self._add_dedication_page()
        
        # Add acknowledgement page
        self.doc.add_page_break()
        self._add_acknowledgement_page()
        
        # Add abstract
        self.doc.add_page_break()
        self._add_abstract(title, abstract)
        
        # Add table of contents
        self.doc.add_page_break()
        self._add_table_of_contents()

    def _add_uoj_approval_page(self, author: str, title: str):
        """Add UoJ approval page"""
        heading = self.doc.add_heading("Certification", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        text = (
            f"The undersigned certify that they have read and hereby recommend for acceptance "
            f"a thesis titled \"{title}\" submitted by {author} in partial fulfillment of "
            f"the requirements for the degree of Doctor of Philosophy in the "
            f"University of Juba."
        )
        self.doc.add_paragraph(text)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Supervisor section
        para = self.doc.add_paragraph("Supervisor:")
        self.doc.add_paragraph("_" * 50)
        self.doc.add_paragraph("[Supervisor Name]")
        self.doc.add_paragraph()
        
        # Date section
        para = self.doc.add_paragraph("Date:")
        self.doc.add_paragraph("_" * 50)

    def _add_uoj_declaration_page(self, author: str):
        """Add UoJ declaration page"""
        heading = self.doc.add_heading("Declaration", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        declaration_text = (
            f"I, {author}, hereby declare that this thesis is my own original work. "
            f"It has not been presented to any other university for the award of a degree. "
            f"All sources used in this thesis have been properly acknowledged."
        )
        self.doc.add_paragraph(declaration_text)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Signature line
        self.doc.add_paragraph("_" * 50)
        self.doc.add_paragraph(f"{author}")
        
        self.doc.add_paragraph()
        
        # Date line
        self.doc.add_paragraph("_" * 50)
        self.doc.add_paragraph("Date")

    def generate_from_topic(self, topic_data: Dict[str, Any]) -> Optional[str]:
        """
        Generate thesis from just a topic/objectives

        Args:
            topic_data: Dictionary containing:
                - title: Thesis title
                - author_name: Author name
                - supervisor: Supervisor name
                - topic: Thesis topic/objectives
                - objectives: List of objectives (optional)

        Returns:
            Path to generated .docx file
        """
        # Create placeholder chapters from objectives
        chapters = {}
        
        objectives = topic_data.get("objectives", [])
        if isinstance(objectives, list):
            for i, objective in enumerate(objectives, 1):
                if i <= 6:  # UoJ has 6 chapters max
                    chapters[i] = {
                        "title": f"Chapter {i}: {objective if isinstance(objective, str) else f'Objective {i}'}",
                        "content": "[Chapter content to be written]\n\n"
                                 f"This chapter addresses the objective: {objective if isinstance(objective, str) else 'See outline'}"
                    }
        
        # If no objectives, create placeholder chapters
        if not chapters:
            for i in range(1, 7):
                chapters[i] = {
                    "title": f"Chapter {i}",
                    "content": "[Chapter content to be written]"
                }

        # Prepare complete thesis input
        thesis_input = {
            "title": topic_data.get("title", "Untitled Thesis"),
            "author_name": topic_data.get("author_name", "Author Name"),
            "supervisor": topic_data.get("supervisor", "Supervisor Name"),
            "topic": topic_data.get("topic", ""),
            "chapters": chapters,
            "abstract": topic_data.get("abstract", ""),
        }

        # Generate complete thesis
        return self.generate_complete_thesis(thesis_input)

    def get_summary(self) -> Dict[str, Any]:
        """Get thesis generation summary"""
        return {
            "university": self.university_name,
            "type": "phd",
            "chapters": len(self.metadata.get("chapters", {})),
            "metadata": self.metadata,
        }
