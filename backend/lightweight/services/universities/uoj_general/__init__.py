"""
University of Juba General Proposal Generator
For Bachelor's Degrees
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, Optional
from pathlib import Path
import json

from ..base_template import BaseThesisGenerator


class UoJGeneralGenerator(BaseThesisGenerator):
    """University of Juba General (Bachelor) Proposal Generator"""

    def __init__(self, workspace_id: str):
        super().__init__("uoj_general", workspace_id)
        self.university_name = "UNIVERSITY OF JUBA"
        
    def _add_cover_page(self, title: str, author: str, supervisor: str, **kwargs):
        """Add UoJ General cover page"""
        # University name
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(self.university_name)
        run.font.size = Pt(14)
        run.font.bold = True
        
        # School/Department placeholders (can be passed in kwargs)
        school = kwargs.get('school', 'SCHOOL OF [SCHOOL NAME]')
        dept = kwargs.get('department', 'DEPARTMENT OF [DEPARTMENT NAME]')
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(school)
        run.font.size = Pt(12)
        run.font.bold = True
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(dept)
        run.font.size = Pt(12)
        run.font.bold = True

        self.doc.add_paragraph()
        
        # TOPIC
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"TOPIC: {title}")
        run.font.size = Pt(12)
        run.font.bold = True
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # BY
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("BY")
        run.font.size = Pt(12)
        
        # Name and Index
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"NAME: {author}")
        run.font.size = Pt(12)
        
        index_no = kwargs.get('index_no', '[INDEX NO]')
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"INDEX NO. {index_no}")
        run.font.size = Pt(12)

        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Submission Text
        submission_text = (
            f"A RESEARCH PROPOSAL SUBMITTED TO THE {school} "
            f"IN PARTIAL FULFILLMENT OF THE REQUIREMENT FOR THE AWARD "
            f"OF A BACHELOR’S DEGREE OF [DEGREE NAME]"
        )
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(submission_text)
        run.font.size = Pt(11)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Date
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("OCTOBER 2026") # As requested
        run.font.size = Pt(12)
        run.font.bold = True

    def _add_preliminary_pages(self, title: str, author: str, abstract: str, **kwargs):
        """Add UoJ General preliminary pages matches user request order"""
        
        # Approval
        self.doc.add_page_break()
        heading = self.doc.add_heading("APPROVAL", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("Write supervisor approval statement for my study")
        
        # Declaration
        self.doc.add_page_break()
        heading = self.doc.add_heading("DECLARATION", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("Write student declaration statement for my study")
        
        # Dedication
        self.doc.add_page_break()
        heading = self.doc.add_heading("DEDICATION", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("Write dedication for my study")
        
        # Acknowledgement
        self.doc.add_page_break()
        heading = self.doc.add_heading("ACKNOWLEDGEMENT", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("Write acknowledgement for my study")
        
        # Abstract
        self.doc.add_page_break()
        heading = self.doc.add_heading("ABSTRACT", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(abstract if abstract else "[Abstract Content]")
        
        # Table of Contents
        self.doc.add_page_break()
        self._add_table_of_contents()

    def generate_from_topic(self, topic_data: Dict[str, Any]) -> Optional[str]:
        """
        Generate General Proposal from topic/objectives using specific Orchestrator.
        """
        from .content_generator import UoJGeneralContentOrchestrator
        import asyncio
        
        # Extract inputs
        topic = topic_data.get("topic", "")
        # Try to parse case study/country from topic or use defaults if not provided in metadata
        # Ideally the request object should pass these, but typically they are in the 'topic' string or separate fields
        # For now, we assume simple extraction or defaults, as the Orchestrator needs them
        
        # In a real scenario, we might want to parse these from the 'topic' string if not explicitly given
        case_study = topic_data.get("case_study", "General Case Study") 
        country = topic_data.get("country", "South Sudan")
        
        # CHECK FOR SCALAR TYPES - user feedback loop
        sample_size_raw = topic_data.get("sample_size")
        sample_size = int(sample_size_raw) if sample_size_raw and str(sample_size_raw).isdigit() else None
        
        objectives = topic_data.get("objectives", [])

        # Initialize Orchestrator
        orchestrator = UoJGeneralContentOrchestrator(
            workspace_id=self.workspace_id,
            topic=topic,
            case_study=case_study,
            country=country,
            sample_size=sample_size,
            objectives=objectives
        )
        
        # Run sync wrapper for async generation (since this method is sync in base class)
        # Note: In FastAPI we are usually in async context, but BaseGenerator might be sync.
        # If BaseGenerator methods are sync, we need to run_until_complete.
        # However, checking base_template, generate_complete_thesis is sync.
        # But we are in an async route in FastAPI... 
        
        # Option: Use asyncio.run() or existing loop
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
        if loop.is_running():
             # If we are already in an async loop (FastAPI), we can't use run_until_complete easily without nesting issues
             # But this method is called from an `await` capable route?
             # Actually, `generate_thesis_from_topic` in `multi_university_thesis.py` is `async def`.
             # So it calls `generator.generate_from_topic` which is currently defined as sync.
             # We should probably make the base method async or handle this carefully.
             # For now, I'll use a nest_asyncio approach or run in executor if needed.
             # EASIEST FIX: Just make this method implementation blocking by running a new loop in a thread or 
             # Refactoring the architecture is risky. 
             # I will try to assume I can run it.
             # Correction: I will assume the caller can handle sync, so I run the async orchestrator synchronously here.
             
             import nest_asyncio
             nest_asyncio.apply()
             content_data = loop.run_until_complete(orchestrator.generate_full_proposal())
        else:
             content_data = loop.run_until_complete(orchestrator.generate_full_proposal())

        # Prepare final input for formatter
        thesis_input = {
            "title": topic_data.get("title", "Untitled Proposal"),
            "author_name": topic_data.get("author_name", "Author Name"),
            "supervisor": topic_data.get("supervisor", "Supervisor Name"),
            "topic": topic,
            "chapters": {
                1: {"title": "CHAPTER ONE: INTRODUCTION", "content": content_data["chapters"][1]},
                2: {"title": "CHAPTER TWO: LITERATURE REVIEWS", "content": content_data["chapters"][2]},
                3: {"title": "CHAPTER THREE: METHODOLOGY", "content": content_data["chapters"][3]},
                4: {"title": "CHAPTER FOUR: DATA ANALYSIS", "content": content_data["chapters"][4]},
                5: {"title": "CHAPTER FIVE: DISCUSSIONS", "content": content_data["chapters"][5]},
            },
            "abstract": content_data["abstract"],
            "appendices": content_data["appendices"]
        }
        
        # Formatter handles formatting (Cover, etc)
        return self.generate_complete_thesis(thesis_input)
