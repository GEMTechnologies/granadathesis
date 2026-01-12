"""
Enhanced DOCX Converter - Well-formatted academic documents

Converts markdown to DOCX with:
- Times New Roman 12pt font
- 1.5 line spacing
- Justified text
- Proper heading styles
- Lists, tables, code blocks, images
- Well organized structure
- Embedded images (downloads and embeds images instead of links)
- Math equations (LaTeX to OMML conversion)
"""
import re
import os
import httpx
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from typing import List, Optional
from io import BytesIO
from PIL import Image

# Math equation support - using math2docx for proper Word equations
try:
    from math2docx import add_math
    MATH2DOCX_AVAILABLE = True
except ImportError:
    MATH2DOCX_AVAILABLE = False
    print("⚠️ math2docx not installed. Falling back to basic math rendering.")

# Fallback imports
try:
    from latex2mathml.converter import convert as latex_to_mathml
    MATH_AVAILABLE = True
except ImportError:
    MATH_AVAILABLE = False

try:
    from lxml import etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False



class EnhancedDOCXConverter:
    """Convert markdown to well-formatted DOCX documents."""
    
    def __init__(self, workspace_id: str = "default"):
        self.doc = None
        self.normal_style = None
        self.workspace_id = workspace_id
        
        # Try multiple paths to find workspace
        candidate_paths = [
            Path(__file__).parent.parent.parent / "thesis_data" / workspace_id,
            Path(__file__).parent.parent.parent.parent / "thesis_data" / workspace_id,
            Path(".") / "thesis_data" / workspace_id,
        ]
        
        self.workspace_dir = None
        for path in candidate_paths:
            if path.exists():
                self.workspace_dir = path
                print(f"✓ Found workspace at: {self.workspace_dir}")
                break
        
        if not self.workspace_dir:
            # Use first candidate even if doesn't exist (for file creation)
            self.workspace_dir = candidate_paths[0]
            print(f"⚠️ Workspace not found at common paths. Using: {self.workspace_dir}")
        
    def convert(self, markdown_content: str, filename: str = "document") -> str:
        """
        Convert markdown to DOCX file.
        
        Returns:
            Path to temporary DOCX file
        """
        # Create document
        self.doc = Document()
        
        # Configure document styles
        self._setup_styles()
        
        # Parse and convert markdown
        self._parse_markdown(markdown_content)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        self.doc.save(temp_file.name)
        return temp_file.name
    
    def _setup_styles(self):
        """Configure document styles for academic formatting."""
        # Set margins (1 inch all around)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Configure Normal style
        self.normal_style = self.doc.styles['Normal']
        font = self.normal_style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        paragraph_format = self.normal_style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = Pt(6)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Configure heading styles
        for level in range(1, 10):
            try:
                heading_style = self.doc.styles[f'Heading {level}']
                heading_font = heading_style.font
                heading_font.name = 'Times New Roman'
                heading_font.bold = True
                
                heading_format = heading_style.paragraph_format
                heading_format.space_before = Pt(12) if level == 1 else Pt(6)
                heading_format.space_after = Pt(6)
                heading_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Set heading font sizes
                if level == 1:
                    heading_font.size = Pt(16)
                elif level == 2:
                    heading_font.size = Pt(14)
                else:
                    heading_font.size = Pt(12)
            except:
                pass
    
    def _parse_markdown(self, content: str):
        """Parse markdown content and convert to DOCX elements."""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Empty line
            if not line.strip():
                i += 1
                continue
            
            # Headings
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                # Strip markdown emphasis markers from headings (headings are bold by style).
                heading_text = re.sub(r'^\*{1,3}(.+?)\*{1,3}$', r'\1', heading_text)
                heading = self.doc.add_heading(heading_text, level=min(level, 9))
                
                # Detect References section - all following paragraphs get hanging indent
                if 'REFERENCES' in heading_text.upper():
                    self.in_references_section = True
                elif level <= 2 and 'REFERENCES' not in heading_text.upper():
                    # Reset when we hit a new major section (level 1 or 2)
                    self.in_references_section = False
                
                # Center only cover page elements; keep headings left-aligned.
                centered_keywords = [
                    'UNIVERSITY OF JUBA',
                    'DECLARATION',
                    'APPROVAL',
                    'DEDICATION',
                    'ACKNOWLEDGEMENTS',
                    'ABSTRACT',
                    'TABLE OF CONTENTS',
                    'LIST OF TABLES',
                    'LIST OF FIGURES',
                    'LIST OF ABBREVIATIONS',
                ]
                if any(kw in heading_text.upper() for kw in centered_keywords):
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Make heading text black and bold
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                    run.bold = True
                
                i += 1
                continue
            
            # Code blocks - check for ASCII diagrams first
            if line.strip().startswith('```'):
                code_lines = []
                i += 1
                language = line.strip()[3:].strip()
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1
                
                code_content = '\n'.join(code_lines)
                
                # Check if this is an ASCII diagram (boxes, arrows, framework)
                if self._is_ascii_diagram(code_content):
                    self._add_visual_framework(code_content)
                else:
                    self._add_code_block(code_content, language)
                continue
            
            # Unordered lists
            if re.match(r'^\s*[-*+]\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^\s*[-*+]\s+', lines[i]):
                    item_text = re.sub(r'^\s*[-*+]\s+', '', lines[i]).strip()
                    list_items.append(item_text)
                    i += 1
                self._add_unordered_list(list_items)
                continue
            
            # Ordered lists
            if re.match(r'^\s*\d+\.\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                    item_text = re.sub(r'^\s*\d+\.\s+', '', lines[i]).strip()
                    list_items.append(item_text)
                    i += 1
                self._add_ordered_list(list_items)
                continue
            
            # Blockquotes
            if line.strip().startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_text = lines[i].strip()[1:].strip()
                    quote_lines.append(quote_text)
                    i += 1
                self._add_blockquote('\n'.join(quote_lines))
                continue
            
            # Tables (basic detection)
            if '|' in line and line.count('|') >= 2:
                # Check if this is a separator row - if so, skip it
                if re.match(r'^\s*\|[\s\-:]+\|\s*$', line):
                    i += 1
                    continue
                
                table_lines = []
                while i < len(lines) and '|' in lines[i] and lines[i].count('|') >= 2:
                    current_line = lines[i].strip()
                    # Skip separator rows completely - various patterns
                    is_separator = (re.match(r'^\|[\s\-:]+\|$', current_line) or 
                                   re.match(r'^[\-:| ]+$', current_line) or
                                   all(c in '-: |' for c in current_line))
                    if not is_separator:
                        table_lines.append(lines[i])
                    i += 1
                if table_lines:
                    self._add_table(table_lines)
                continue
            
            # HTML anchor tags for bookmarks: <a id="ref_..."></a>
            anchor_match = re.match(r'<a\s+id=["\']([^"\']+)["\']\s*>\s*</a>', line.strip())
            if anchor_match:
                bookmark_name = anchor_match.group(1)
                self._add_bookmark(bookmark_name)
                i += 1
                continue
            
            # Markdown images: ![alt text](url)
            if line.strip().startswith('!['):
                match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line.strip())
                if match:
                    alt_text = match.group(1)
                    image_url = match.group(2)
                    # Create a paragraph and embed the image
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the image
                    self._add_image(p, image_url, alt_text)
                    i += 1
                    continue
            
            # Block math equations: $$...$$ or \[...\]
            if line.strip().startswith('$$') or line.strip().startswith('\\['):
                is_bracket = line.strip().startswith('\\[')
                start_delim = '\\[' if is_bracket else '$$'
                end_delim = '\\]' if is_bracket else '$$'
                
                # Check if it's a single-line block math
                if is_bracket:
                    # \[equation\] format
                    if line.strip().endswith('\\]') and len(line.strip()) > 4:
                        latex = line.strip()[2:-2].strip()
                        self._add_math_block(latex)
                        i += 1
                        continue
                else:
                    # $$equation$$ format
                    if line.strip().endswith('$$') and len(line.strip()) > 4:
                        latex = line.strip()[2:-2].strip()
                        self._add_math_block(latex)
                        i += 1
                        continue
                
                # Multi-line block math
                math_lines = []
                i += 1
                while i < len(lines):
                    current = lines[i].strip()
                    if (is_bracket and current.endswith('\\]')) or (not is_bracket and current.endswith('$$')):
                        # Check if there's content before the closing delimiter
                        if is_bracket and current != '\\]':
                            math_lines.append(current[:-2])
                        elif not is_bracket and current != '$$':
                            math_lines.append(current[:-2])
                        i += 1
                        break
                    math_lines.append(lines[i])
                    i += 1
                
                latex = '\n'.join(math_lines).strip()
                if latex:
                    self._add_math_block(latex)
                continue
            
            # DETECT UNDELIMITED STANDALONE LATEX EQUATIONS
            # Lines that look like: r = \frac{...}{...} or Y = \beta_0 + ...
            stripped = line.strip()
            if self._is_standalone_latex_equation(stripped):
                # Extract the LaTeX and render as block equation
                latex = self._extract_latex_from_line(stripped)
                if latex:
                    self._add_math_block(latex)
                    i += 1
                    continue
            
            # Regular paragraph (may contain inline math)
            paragraph_text = line.strip()
            if paragraph_text:
                self._add_paragraph(paragraph_text)
            i += 1

    def _latex_to_omml(self, latex: str) -> Optional[OxmlElement]:
        """Convert LaTeX to Office Math Markup Language (OMML) XML element."""
        if not MATH_AVAILABLE:
            return None
        
        try:
            # Convert LaTeX to MathML (for validation, though we don't use it directly)
            mathml = latex_to_mathml(latex)
            
            # Create OMML elements - python-docx OxmlElement handles namespaces
            # Create an oMathPara for block equations or oMath for inline
            omath = OxmlElement('m:oMath')
            
            # Add the equation text as a math run
            r = OxmlElement('m:r')
            
            # Math text element
            t = OxmlElement('m:t')
            t.text = latex
            
            r.append(t)
            omath.append(r)
            
            return omath
            
        except Exception as e:
            print(f"⚠️ LaTeX-to-OMML conversion failed: {e}")
            return None
    
    def _add_math_block(self, latex: str):
        """Add a block math equation to the document using math2docx."""
        # Sanitize LaTeX for XML compatibility
        def sanitize_for_xml(s: str) -> str:
            """Remove or replace characters that are invalid in XML."""
            # Remove NULL bytes and control characters (except tab, newline, carriage return)
            import re
            # Keep only valid XML characters
            return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)
        
        def fix_latex_for_math2docx(s: str) -> str:
            """Fix LaTeX constructs that math2docx doesn't handle well."""
            # Replace \bar{x} with \overline{x} - math2docx has a bug with \bar
            result = re.sub(r'\\bar\{([^}]+)\}', r'\\overline{\1}', s)
            # Replace \hat{x} with proper form if needed
            # (add more fixes as discovered)
            return result
        
        sanitized_latex = sanitize_for_xml(latex)
        fixed_latex = fix_latex_for_math2docx(sanitized_latex)
        
        if MATH2DOCX_AVAILABLE:
            try:
                # Use math2docx - it handles everything including paragraph creation
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_math(p, fixed_latex)
                return
            except Exception as e:
                print(f"⚠️ math2docx failed: {e}")
        
        # Fallback to basic OMML (may not render perfectly)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if MATH_AVAILABLE and LXML_AVAILABLE:
            try:
                omml = self._latex_to_omml(sanitized_latex)
                if omml is not None:
                    p._element.append(omml)
                    return
            except Exception as e:
                print(f"⚠️ Failed to add OMML math: {e}")
        
        # Ultimate fallback: Styled text (make it readable)
        # Convert LaTeX to readable format for display
        readable = sanitized_latex
        readable = readable.replace('\\frac{', '(')
        readable = readable.replace('}{', ')/(')
        readable = readable.replace('\\sqrt{', '√(')
        readable = readable.replace('\\sum', 'Σ')
        readable = readable.replace('\\bar{', '')
        readable = readable.replace('\\alpha', 'α')
        readable = readable.replace('\\beta', 'β')
        readable = readable.replace('\\gamma', 'γ')
        readable = readable.replace('\\delta', 'δ')
        readable = readable.replace('\\epsilon', 'ε')
        readable = readable.replace('\\varepsilon', 'ε')
        readable = readable.replace('\\mu', 'μ')
        readable = readable.replace('\\sigma', 'σ')
        readable = readable.replace('\\chi', 'χ')
        readable = readable.replace('\\rho', 'ρ')
        readable = readable.replace('\\theta', 'θ')
        readable = readable.replace('\\lambda', 'λ')
        readable = readable.replace('\\omega', 'ω')
        readable = readable.replace('\\infty', '∞')
        readable = readable.replace('\\pm', '±')
        readable = readable.replace('\\times', '×')
        readable = readable.replace('\\cdot', '·')
        readable = readable.replace('\\leq', '≤')
        readable = readable.replace('\\geq', '≥')
        readable = readable.replace('\\neq', '≠')
        readable = readable.replace('\\approx', '≈')
        readable = readable.replace('{', '')
        readable = readable.replace('}', '')
        readable = readable.replace('\\', '')
        
        run = p.add_run(readable)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(12)
        run.italic = True
    
    def _add_inline_math(self, paragraph, latex: str):
        """Add inline math equation to a paragraph using math2docx."""
        # Sanitize LaTeX for XML compatibility
        def sanitize_for_xml(s: str) -> str:
            """Remove or replace characters that are invalid in XML."""
            import re
            return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)
        
        def fix_latex_for_math2docx(s: str) -> str:
            """Fix LaTeX constructs that math2docx doesn't handle well."""
            import re
            # Replace \bar{x} with \overline{x} - math2docx has a bug with \bar
            return re.sub(r'\\bar\{([^}]+)\}', r'\\overline{\1}', s)
        
        sanitized_latex = sanitize_for_xml(latex)
        fixed_latex = fix_latex_for_math2docx(sanitized_latex)
        
        if MATH2DOCX_AVAILABLE:
            try:
                add_math(paragraph, fixed_latex)
                return
            except Exception as e:
                print(f"⚠️ math2docx inline failed: {e}")
        
        # Fallback to basic OMML
        if MATH_AVAILABLE and LXML_AVAILABLE:
            try:
                omml = self._latex_to_omml(fixed_latex)
                if omml is not None:
                    paragraph._element.append(omml)
                    return
            except Exception as e:
                print(f"⚠️ Failed to add inline OMML math: {e}")
        
        # Ultimate fallback: Convert LaTeX to readable Unicode
        readable = fixed_latex
        readable = readable.replace('\\frac{', '(')
        readable = readable.replace('}{', ')/(')
        readable = readable.replace('\\sqrt{', '√(')
        readable = readable.replace('\\sum', 'Σ')
        readable = readable.replace('\\bar{', '')
        readable = readable.replace('\\overline{', '')
        readable = readable.replace('\\alpha', 'α')
        readable = readable.replace('\\beta', 'β')
        readable = readable.replace('\\gamma', 'γ')
        readable = readable.replace('\\delta', 'δ')
        readable = readable.replace('\\epsilon', 'ε')
        readable = readable.replace('\\varepsilon', 'ε')
        readable = readable.replace('\\mu', 'μ')
        readable = readable.replace('\\sigma', 'σ')
        readable = readable.replace('\\chi', 'χ')
        readable = readable.replace('\\rho', 'ρ')
        readable = readable.replace('\\theta', 'θ')
        readable = readable.replace('\\lambda', 'λ')
        readable = readable.replace('\\omega', 'ω')
        readable = readable.replace('\\infty', '∞')
        readable = readable.replace('\\pm', '±')
        readable = readable.replace('\\times', '×')
        readable = readable.replace('\\cdot', '·')
        readable = readable.replace('\\leq', '≤')
        readable = readable.replace('\\geq', '≥')
        readable = readable.replace('\\neq', '≠')
        readable = readable.replace('\\approx', '≈')
        readable = readable.replace('{', '')
        readable = readable.replace('}', '')
        readable = readable.replace('\\', '')
        
        run = paragraph.add_run(readable)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(12)
        run.italic = True

    def _is_standalone_latex_equation(self, line: str) -> bool:
        """Check if a line is a standalone LaTeX equation (not wrapped in delimiters)."""
        # Skip if already has math delimiters
        if line.startswith('$') or line.startswith('\\[') or line.startswith('\\('):
            return False
        
        # Check for common LaTeX math commands that indicate an equation
        latex_indicators = [
            r'\\frac\{',
            r'\\sqrt\{',
            r'\\sum',
            r'\\prod',
            r'\\int',
            r'\\bar\{',
            r'\\hat\{',
            r'\\tilde\{',
            r'\\vec\{',
            r'\\alpha',
            r'\\beta',
            r'\\gamma',
            r'\\delta',
            r'\\epsilon',
            r'\\varepsilon',
            r'\\mu',
            r'\\sigma',
            r'\\chi',
            r'\\rho',
            r'\\theta',
            r'\\lambda',
            r'\\omega',
            r'\\infty',
            r'\\partial',
            r'\\nabla',
            r'\\pm',
            r'\\cdot',
            r'\\times',
            r'\\leq',
            r'\\geq',
            r'\\neq',
            r'\\approx',
        ]
        
        # Check if line contains LaTeX commands
        for indicator in latex_indicators:
            if re.search(indicator, line):
                # Additional check: should look like an equation (contains = or is mostly math)
                # e.g., "r = \frac{...}" or just "\frac{...}{...}"
                if '=' in line or line.startswith('\\'):
                    return True
                # Check if it's a significant portion of the line
                non_text_ratio = len(re.findall(r'\\[a-zA-Z]+', line)) / max(len(line.split()), 1)
                if non_text_ratio > 0.3:
                    return True
        
        return False

    def _extract_latex_from_line(self, line: str) -> str:
        """Extract LaTeX content from a line, removing any surrounding text if minimal."""
        # If it's primarily an equation (like "r = \frac{...}"), return the whole thing
        # Otherwise try to extract just the equation part
        
        # Pattern: starts with optional variable = then LaTeX
        match = re.match(r'^([A-Za-z]+\s*=\s*)?(\\.*?)$', line)
        if match:
            return line.strip()
        
        # Just return the line as-is for the equation
        return line.strip()

    def _add_paragraph(self, text: str):
        """Add a justified paragraph with formatting."""
        # Check if we're in references section (set by heading detection)
        if hasattr(self, 'in_references_section') and self.in_references_section:
            self._add_reference_paragraph(text)
            return
            
        # Parse inline formatting (bold, italic, links)
        p = self.doc.add_paragraph()
        
        # Auto-centering for cover page elements
        centered_patterns = [
            r'^\*\*BY\*\*$', 
            r'^---+$', 
            r'^DOCTOR OF PHILOSOPHY$', 
            r'^University of Juba$',
            r'^\*\*University of Juba\*\*$',
            r'^[A-Z][a-z]+ \d{4}$' # Date like November 2026
        ]
        
        is_centered = any(re.match(pattern, text.strip()) for pattern in centered_patterns)
        # Do not auto-center short bold lines; reserve centering for explicit patterns above.
            
        if is_centered:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        p.style = self.normal_style
        
        # Parse inline markdown
        self._add_formatted_text(p, text)
    
    def _add_reference_paragraph(self, text: str):
        """Add a reference paragraph with APA 7 hanging indent formatting."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # APA 7 Hanging indent: first line flush left, subsequent lines indented 0.5"
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)  # 12pt spacing between references
        
        # Add text with formatting
        self._add_formatted_text(p, text)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, links, code, math, sub/sup)."""
        # First, normalize LaTeX delimiters: convert \(...\) to $...$
        # This handles both formats uniformly
        normalized_text = re.sub(r'\\\((.+?)\\\)', r'$\1$', text)
        # Also convert \[...\] to $$...$$ for display math
        normalized_text = re.sub(r'\\\[(.+?)\\\]', r'$$\1$$', normalized_text)
        
        # DETECT AND WRAP UNDELIMITED LATEX EQUATIONS
        # Pattern to find LaTeX commands that aren't already wrapped in $..$ or $$..$$
        # These are common LaTeX math patterns that should be rendered as equations
        latex_commands = [
            r'\\frac\{',      # Fractions
            r'\\sum',          # Summation
            r'\\bar\{',        # Bar notation
            r'\\sqrt\{',       # Square root
            r'\\alpha',        # Greek letters
            r'\\beta',
            r'\\gamma',
            r'\\delta',
            r'\\epsilon',
            r'\\varepsilon',
            r'\\mu',
            r'\\sigma',
            r'\\chi',
            r'\\rho',
            r'\\int',          # Integral
            r'\\prod',         # Product
            r'\\lim',          # Limit
            r'\\infty',        # Infinity
            r'\\partial',      # Partial derivative
            r'\\nabla',        # Nabla
            r'\\times',        # Multiplication
            r'\\div',          # Division
            r'\\pm',           # Plus-minus
            r'\\leq',          # Less than or equal
            r'\\geq',          # Greater than or equal
            r'\\neq',          # Not equal
            r'\\approx',       # Approximately
            r'\^{',            # Superscript braces
            r'_{',             # Subscript braces
        ]
        
        # Check if text contains undelimited LaTeX (not already in $ or $$)
        def wrap_undelimited_latex(text):
            """Find and wrap undelimited LaTeX equations."""
            # Pattern to find potential LaTeX expressions not in delimiters
            # Match sequences containing LaTeX commands followed by content
            # This regex finds expressions like: r = \frac{...}{...}
            
            # Match common equation patterns: variable = \latex_command{...}
            patterns = [
                # Pattern: variable = \frac{...}{...} (equations with fractions)
                (r'([A-Za-z])\s*=\s*(\\frac\{[^}]+\}\{[^}]+\})', r'$\1 = \2$'),
                # Pattern: standalone \frac{...}{...} followed by content until newline or period
                (r'(\\frac\{[^}]+\}\{\\sqrt\{[^}]+\}\})', r'$\1$'),
                # Pattern: Y = β₀ + β₁X₁ + ... (regression equations with Greek)
                (r'([A-Z])\s*=\s*(\\beta[_0-9]*\s*(?:\+\s*\\beta[_0-9]*\s*[A-Z][_0-9]*\s*)+(?:\+\s*\\varepsilon)?)', r'$\1 = \2$'),
                # General: Any line starting with backslash commands and containing = 
                (r'(?<!\$)([A-Za-z]+\s*=\s*[^$\n]*\\(?:frac|sum|bar|sqrt|alpha|beta|gamma|int)[^$\n]*?)(?=\s*\n|\s*$|(?:\.\s))', r'$\1$'),
            ]
            
            result = text
            for pattern, replacement in patterns:
                result = re.sub(pattern, replacement, result)
            
            return result
        
        # Apply undelimited LaTeX wrapping
        normalized_text = wrap_undelimited_latex(normalized_text)
        
        # Split by markdown patterns including inline math $...$ AND HTML sub/sup tags
        # Note: We use (?<!\$) and (?!\$) to avoid matching $$ (block math)
        parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\)|(?<!\$)\$(?!\$).*?(?<!\$)\$(?!\$)|<sub>.*?</sub>|<sup>.*?</sup>)', normalized_text)
        
        for part in parts:
            if not part:
                continue
            
            # HTML subscript: <sub>...</sub>
            if part.startswith('<sub>') and part.endswith('</sub>'):
                sub_text = part[5:-6]  # Extract text between tags
                run = paragraph.add_run(sub_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.subscript = True
            # HTML superscript: <sup>...</sup>
            elif part.startswith('<sup>') and part.endswith('</sup>'):
                sup_text = part[5:-6]  # Extract text between tags
                run = paragraph.add_run(sup_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.superscript = True
            # Inline math: $...$
            elif part.startswith('$') and part.endswith('$') and not part.startswith('$$'):
                latex = part[1:-1].strip()
                if latex:
                    self._add_inline_math(paragraph, latex)
            # Bold + italic
            elif part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            # Bold
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            # Italic
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            # Inline code
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(11)
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            # Links - create REAL clickable hyperlinks
            elif part.startswith('[') and '](' in part:
                match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if match:
                    link_text = match.group(1)
                    link_url = match.group(2)
                    
                    # Create a real hyperlink using relationship
                    self._add_hyperlink(paragraph, link_text, link_url)
            else:
                # Regular text
                run = paragraph.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def _add_hyperlink(self, paragraph, text: str, url: str):
        """Add a hyperlink - internal bookmark for #ref_ links, external for URLs."""
        
        # Check if this is an internal reference link (starts with #ref_)
        is_internal = url.startswith('#ref_') or url.startswith('#')
        
        if is_internal:
            # Internal bookmark link - link to reference in same document
            bookmark_name = url.lstrip('#')
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), bookmark_name)
        else:
            # External URL link
            r_id = paragraph.part.relate_to(
                url,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            )
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
        
        # Create run inside hyperlink
        new_run = OxmlElement('w:r')
        
        # Run properties
        rPr = OxmlElement('w:rPr')
        
        # Font name
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        
        # Font size (12pt = 24 half-points)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        
        # BLACK color (not blue) - user requested all text black
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')  # Black
        rPr.append(color)
        
        # NO underline for cleaner look
        # (removed underline element)
        
        new_run.append(rPr)
        
        # Add text
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def _add_bookmark(self, bookmark_name: str):
        """Create a Word bookmark for internal linking (used for references)."""
        # Create an invisible paragraph with a bookmark
        p = self.doc.add_paragraph()
        
        # Create bookmark start element  
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 999999))
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        # Create bookmark end element
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 999999))
        
        # Insert bookmark into paragraph
        p._p.insert(0, bookmark_start)
        p._p.append(bookmark_end)
        
        # Make paragraph invisible (no text, zero spacing)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    
    def _add_unordered_list(self, items: List[str]):
        """Add unordered list with proper formatting."""
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Format list item
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def _add_ordered_list(self, items: List[str]):
        """Add ordered list with proper formatting."""
        for item in items:
            p = self.doc.add_paragraph(item, style='List Number')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Format list item
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def _add_blockquote(self, text: str):
        """Add blockquote with proper formatting."""
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.style = self.normal_style
        
        # Add left border effect (indentation)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        
        # Italic text
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
    
    def _is_ascii_diagram(self, content: str) -> bool:
        """Detect if content is an ASCII diagram (boxes, arrows, framework)."""
        # Common ASCII diagram patterns
        patterns = [
            r'\+[-=]+\+',           # Box borders: +---+
            r'\|.*\|',              # Box sides: |text|
            r'[─═│┌┐└┘├┤┬┴┼]',      # Unicode box characters
            r'-->|<--|->|<-|→|←',   # Arrows
            r'\[.*\].*[-=]+.*\[',   # [Box]---[Box]
            r'╔|╗|╚|╝|║|═',         # Double-line box chars
        ]
        
        lines = content.strip().split('\n')
        
        # Check for multiple box-like patterns
        box_indicators = 0
        for line in lines:
            for pattern in patterns:
                if re.search(pattern, line):
                    box_indicators += 1
                    break
        
        # If more than 3 lines have box patterns, it's likely a diagram
        return box_indicators >= 3
    
    def _add_visual_framework(self, ascii_content: str):
        """Convert ASCII diagram to AI-generated professional image and embed it."""
        import asyncio
        
        # Parse ASCII to extract components for better prompt
        lines = ascii_content.strip().split('\n')
        
        # Extract text from boxes (content between | | or [ ])
        components = []
        for line in lines:
            box_texts = re.findall(r'\|\s*([^|]+?)\s*\|', line)
            bracket_texts = re.findall(r'\[\s*([^\]]+?)\s*\]', line)
            components.extend(box_texts + bracket_texts)
        
        # Remove duplicates and empty strings
        components = [c.strip() for c in components if c.strip()]
        components = list(dict.fromkeys(components))
        
        if not components:
            # Fallback: just show as formatted text
            self._add_code_block(ascii_content, "")
            return
        
        # FAST: Table-based framework (images pre-generated during writing phase)
        component_list = ", ".join(components[:6])
        
        # Create table-based framework visualization
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        num_rows = (len(components) + 1) // 2
        if num_rows == 0:
            num_rows = 1
        
        table = self.doc.add_table(rows=num_rows, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for idx, component in enumerate(components):
            row_idx = idx // 2
            col_idx = idx % 2
            
            if row_idx < len(table.rows):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = component
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.bold = True
                
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E8F4FD')
                cell._tc.get_or_add_tcPr().append(shading)
                cell.width = Inches(2.5)
        
        # Add caption
        caption_p = self.doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run(f"Figure: Conceptual Framework - {component_list[:50]}")
        caption_run.font.name = 'Times New Roman'
        caption_run.font.size = Pt(10)
        caption_run.italic = True
        
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
        
        if len(components) > 2:
            arrow_p = self.doc.add_paragraph()
            arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = arrow_p.add_run("↓  →  ↔  (Relationships indicated by arrows)")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
    
    def _add_code_block(self, code: str, language: str = ''):
        """Add code block with monospace font."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add language label if provided
        if language:
            run = p.add_run(f'[{language}]\n')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add code
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        # Background color
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    def _add_table(self, table_lines: List[str]):
        """Add table from markdown table syntax."""
        if not table_lines:
            return
        
        # Parse table rows, filtering out separator rows (---|---|---)
        rows = []
        for line in table_lines:
            stripped = line.strip()
            # Skip markdown table separator rows - catch all variations
            if (re.match(r'^\|[\s\-:]+\|$', stripped) or 
                all(c in '-: |' for c in stripped)):
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Only add non-empty rows
            if cells and any(c for c in cells):
                rows.append(cells)
        
        if not rows:
            return
        
        # Create table
        num_cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Populate table
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    # Format cell text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
        
        # Add spacing after table
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
    
    def _add_image(self, paragraph, image_url: str, alt_text: str = ""):
        """Download and embed an image in the document."""
        image_bytes = None
        resolved_path = None
        
        try:
            # STEP 1: Try to load from local file first
            # Support multiple path patterns
            if image_url.startswith('/') or image_url.startswith('./') or image_url.startswith('../'):
                # Relative or absolute path
                resolved_path = self._resolve_image_path(image_url)
                if resolved_path and resolved_path.exists():
                    image_bytes = resolved_path.read_bytes()
                    print(f"✓ Loaded local image: {resolved_path}")
            
            # STEP 2: Try workspace directory lookup
            if not image_bytes:
                image_bytes = self._load_from_workspace(image_url)
            
            # STEP 3: Try as direct URL (HTTP/HTTPS)
            if not image_bytes and (image_url.startswith('http://') or image_url.startswith('https://')):
                image_bytes = self._download_image(image_url)
            
            # STEP 4: If still no image, try as data URI (base64)
            if not image_bytes and image_url.startswith('data:'):
                image_bytes = self._decode_data_uri(image_url)
            
            # If we got the bytes, embed the image
            if image_bytes:
                self._embed_image_bytes(paragraph, image_bytes, alt_text)
                return
            
            # If all methods failed
            raise ValueError(f"Could not load image from any source: {image_url}")
        
        except Exception as e:
            import traceback
            error_msg = f"Error embedding image {image_url}: {str(e)}"
            print(f"❌ {error_msg}")
            print(traceback.format_exc())
            
            # Fallback: add as text reference with warning
            run = paragraph.add_run(f"\n[⚠️ Image Not Embedded: {alt_text or image_url}]")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red to indicate error
    
    def _resolve_image_path(self, image_url: str) -> Optional[Path]:
        """Resolve a local image path."""
        try:
            # Handle relative paths
            if image_url.startswith('../'):
                # Relative path - resolve from workspace
                path = (self.workspace_dir / image_url).resolve()
            elif image_url.startswith('./'):
                # Current directory relative
                path = (self.workspace_dir / image_url).resolve()
            elif image_url.startswith('/'):
                # Absolute path
                path = Path(image_url)
            else:
                return None
            
            return path if path.exists() else None
        except Exception as e:
            print(f"Path resolution error: {e}")
            return None
    
    def _load_from_workspace(self, image_url: str) -> Optional[bytes]:
        """Try to load image from various workspace locations."""
        try:
            # Extract filename and relative path
            image_path = Path(image_url)
            filename = image_path.name
            if not filename:
                return None
            
            # Try common image storage locations
            search_paths = [
                # First try the path as-is relative to workspace (e.g., figures/bar_gender.png)
                self.workspace_dir / image_url,
                # Then try with chapters prefix (e.g., chapters/figures/bar_gender.png)
                self.workspace_dir / "chapters" / image_url,
                # Then try filename-only searches in common directories
                self.workspace_dir / "chapters" / "figures" / filename,
                self.workspace_dir / "images" / filename,
                self.workspace_dir / "figures" / filename,
                self.workspace_dir / filename,
                self.workspace_dir / "static" / filename,
                self.workspace_dir / "assets" / filename,
                self.workspace_dir / "media" / filename,
            ]
            
            for path in search_paths:
                if path.exists():
                    print(f"✓ Found image at: {path}")
                    return path.read_bytes()
            
            # Log all paths tried for debugging
            print(f"❌ File not found: {image_url}")
            print(f"   Checked: {self.workspace_dir / image_url}")
            return None
        except Exception as e:
            print(f"Workspace lookup error: {e}")
            return None
    
    def _download_image(self, image_url: str) -> Optional[bytes]:
        """Download image from URL with retry logic."""
        max_retries = 3
        timeout = 20.0  # Increased timeout
        
        for attempt in range(max_retries):
            try:
                print(f"Downloading image (attempt {attempt + 1}/{max_retries}): {image_url}")
                
                # Use httpx with more aggressive settings
                with httpx.Client(
                    timeout=timeout,
                    follow_redirects=True,
                    verify=True,
                    limits=httpx.Limits(max_connections=5, max_keepalive_connections=2)
                ) as client:
                    response = client.get(image_url)
                    response.raise_for_status()
                    
                    # Check content type
                    content_type = response.headers.get('content-type', '').lower()
                    if 'image' not in content_type:
                        print(f"Warning: Unexpected content-type: {content_type}")
                    
                    print(f"✓ Downloaded image successfully ({len(response.content)} bytes)")
                    return response.content
            
            except httpx.TimeoutException:
                print(f"Timeout downloading image (attempt {attempt + 1}/{max_retries})")
                if attempt < max_retries - 1:
                    import time
                    time.sleep(1)  # Wait before retry
            except httpx.HTTPStatusError as e:
                print(f"HTTP error {e.response.status_code}: {image_url}")
                break  # Don't retry on HTTP errors
            except Exception as e:
                print(f"Download error: {e}")
                if attempt < max_retries - 1:
                    import time
                    time.sleep(1)  # Wait before retry
        
        return None
    
    def _decode_data_uri(self, data_uri: str) -> Optional[bytes]:
        """Decode base64 data URI."""
        try:
            # Format: data:image/png;base64,iVBORw0KGgoAAAANS...
            if ',' not in data_uri:
                return None
            
            import base64
            header, data = data_uri.split(',', 1)
            if 'base64' not in header:
                return None
            
            return base64.b64decode(data)
        except Exception as e:
            print(f"Data URI decode error: {e}")
            return None
    
    def _embed_image_bytes(self, paragraph, image_bytes: bytes, alt_text: str = ""):
        """Embed image bytes into paragraph."""
        try:
            # Process image with PIL to ensure compatibility
            img = Image.open(BytesIO(image_bytes))
            
            # Convert to RGB if necessary (for PNG with transparency)
            if img.mode in ('RGBA', 'LA', 'P'):
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                rgb_img.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                img = rgb_img
            
            # Save to BytesIO
            img_bytes_io = BytesIO()
            img.save(img_bytes_io, format='PNG')
            img_bytes_io.seek(0)
            
            # Add image to paragraph
            run = paragraph.add_run()
            
            # Calculate size (max width 6 inches, maintain aspect ratio)
            max_width_inches = 6.0
            max_width_px = max_width_inches * 96  # 96 DPI
            
            width, height = img.size
            
            # Scale down if too wide
            if width > max_width_px:
                scale = max_width_px / width
                width = int(max_width_px)
                height = int(height * scale)
            
            # Convert to inches (96 DPI)
            width_inches = Inches(width / 96)
            height_inches = Inches(height / 96)
            
            run.add_picture(img_bytes_io, width=width_inches, height=height_inches)
            
            # Add caption if alt text provided
            if alt_text:
                caption = paragraph.add_run(f"\n{alt_text}")
                caption.font.name = 'Times New Roman'
                caption.font.size = Pt(10)
                caption.font.italic = True
                caption.font.color.rgb = RGBColor(128, 128, 128)
            
            print(f"✓ Image embedded successfully: {width}x{height}px")
        
        except Exception as e:
            print(f"Error embedding image bytes: {e}")
            raise


def convert_markdown_to_docx_enhanced(content: str, filename: str = "document", workspace_id: str = "default") -> str:
    """
    Convert markdown to well-formatted DOCX.
    
    Args:
        content: Markdown content
        filename: Output filename (without extension)
        workspace_id: Workspace ID for resolving local image paths
        
    Returns:
        Path to temporary DOCX file
    """
    converter = EnhancedDOCXConverter(workspace_id=workspace_id)
    return converter.convert(content, filename)
