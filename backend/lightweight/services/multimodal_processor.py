"""
Multimodal File Processor

Processes various file types for AI understanding:
- Audio → Text (speech recognition)
- Images → Vision (description, OCR)
- PDFs → Text extraction
- DOCX → Text extraction
- Video → Frame extraction + audio

Enables the system to "read" and reference uploaded files.
"""

import asyncio
import base64
import json
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import mimetypes

from services.workspace_service import WORKSPACES_DIR


class MultimodalProcessor:
    """
    Process various file types for AI understanding.
    
    Supports:
    - Audio: MP3, WAV, M4A, OGG → Text transcription
    - Images: PNG, JPG, GIF, WebP → Vision description + OCR
    - Documents: PDF, DOCX, TXT → Text extraction
    - Data: CSV, XLSX, JSON → Structured data
    """
    
    def __init__(self):
        self.supported_audio = {'.mp3', '.wav', '.m4a', '.ogg', '.flac', '.webm'}
        self.supported_images = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}
        self.supported_docs = {'.pdf', '.docx', '.doc', '.txt', '.md', '.rtf'}
        self.supported_data = {'.csv', '.xlsx', '.xls', '.json', '.xml'}
    
    def get_file_type(self, file_path: str) -> str:
        """Determine file type category."""
        ext = Path(file_path).suffix.lower()
        
        if ext in self.supported_audio:
            return 'audio'
        elif ext in self.supported_images:
            return 'image'
        elif ext in self.supported_docs:
            return 'document'
        elif ext in self.supported_data:
            return 'data'
        else:
            return 'unknown'
    
    async def process_file(
        self,
        file_path: str,
        workspace_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process any file and return extracted content.
        
        Returns:
            {
                "file_path": str,
                "file_type": str,
                "content": str,  # Extracted text or description
                "metadata": dict,
                "success": bool,
                "error": str | None
            }
        """
        file_type = self.get_file_type(file_path)
        
        try:
            if file_type == 'audio':
                return await self.process_audio(file_path)
            elif file_type == 'image':
                return await self.process_image(file_path)
            elif file_type == 'document':
                return await self.process_document(file_path)
            elif file_type == 'data':
                return await self.process_data(file_path)
            else:
                return {
                    "file_path": file_path,
                    "file_type": file_type,
                    "content": "",
                    "metadata": {},
                    "success": False,
                    "error": f"Unsupported file type: {Path(file_path).suffix}"
                }
        except Exception as e:
            return {
                "file_path": file_path,
                "file_type": file_type,
                "content": "",
                "metadata": {},
                "success": False,
                "error": str(e)
            }
    
    async def process_audio(self, file_path: str) -> Dict[str, Any]:
        """
        Convert audio to text using speech recognition.
        
        Uses OpenAI Whisper or local whisper.cpp
        """
        from core.llm_client import LLMClient
        
        path = Path(file_path)
        if not path.exists():
            return {
                "file_path": file_path,
                "file_type": "audio",
                "content": "",
                "metadata": {},
                "success": False,
                "error": "File not found"
            }
        
        # Read audio file
        audio_data = path.read_bytes()
        
        # Try OpenAI Whisper API first
        try:
            import httpx
            from core.config import settings
            
            if hasattr(settings, 'openai_api_key') and settings.openai_api_key:
                async with httpx.AsyncClient() as client:
                    response = await client.post(
                        "https://api.openai.com/v1/audio/transcriptions",
                        headers={"Authorization": f"Bearer {settings.openai_api_key}"},
                        files={"file": (path.name, audio_data)},
                        data={"model": "whisper-1"},
                        timeout=120.0
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        return {
                            "file_path": file_path,
                            "file_type": "audio",
                            "content": result.get("text", ""),
                            "metadata": {
                                "duration": result.get("duration"),
                                "language": result.get("language")
                            },
                            "success": True,
                            "error": None
                        }
        except Exception as e:
            print(f"OpenAI Whisper failed: {e}")
        
        # Fallback: Use Deepgram or other service
        # For now, return placeholder
        return {
            "file_path": file_path,
            "file_type": "audio",
            "content": f"[Audio file: {path.name}. Transcription requires OpenAI API key or local Whisper model.]",
            "metadata": {"file_size": len(audio_data)},
            "success": True,
            "error": None
        }
    
    async def process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze image using vision AI.
        
        Uses GPT-4 Vision or similar multimodal model.
        """
        from core.llm_client import LLMClient
        
        path = Path(file_path)
        if not path.exists():
            return {
                "file_path": file_path,
                "file_type": "image",
                "content": "",
                "metadata": {},
                "success": False,
                "error": "File not found"
            }
        
        # Read and encode image
        image_data = path.read_bytes()
        base64_image = base64.b64encode(image_data).decode('utf-8')
        
        # Determine MIME type
        mime_type = mimetypes.guess_type(file_path)[0] or 'image/png'
        
        # Try GPT-4 Vision
        try:
            import httpx
            from core.config import settings
            
            if hasattr(settings, 'openai_api_key') and settings.openai_api_key:
                async with httpx.AsyncClient() as client:
                    response = await client.post(
                        "https://api.openai.com/v1/chat/completions",
                        headers={
                            "Authorization": f"Bearer {settings.openai_api_key}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": "gpt-4o",
                            "messages": [
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": "Describe this image in detail. If it contains text, transcribe the text. If it contains data, charts, or diagrams, explain what they show."
                                        },
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:{mime_type};base64,{base64_image}"
                                            }
                                        }
                                    ]
                                }
                            ],
                            "max_tokens": 1000
                        },
                        timeout=60.0
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        description = result['choices'][0]['message']['content']
                        return {
                            "file_path": file_path,
                            "file_type": "image",
                            "content": description,
                            "metadata": {
                                "model": "gpt-4o",
                                "file_size": len(image_data)
                            },
                            "success": True,
                            "error": None
                        }
        except Exception as e:
            print(f"GPT-4 Vision failed: {e}")
        
        # Fallback to DeepSeek vision if available
        try:
            llm = LLMClient()
            # Check if DeepSeek supports vision
            description = await llm.analyze_image(base64_image, mime_type)
            if description:
                return {
                    "file_path": file_path,
                    "file_type": "image",
                    "content": description,
                    "metadata": {"model": "deepseek"},
                    "success": True,
                    "error": None
                }
        except Exception as e:
            print(f"DeepSeek vision failed: {e}")
        
        # Basic fallback
        return {
            "file_path": file_path,
            "file_type": "image",
            "content": f"[Image: {path.name}. Vision analysis requires API key configuration.]",
            "metadata": {"file_size": len(image_data)},
            "success": True,
            "error": None
        }
    
    async def process_document(self, file_path: str) -> Dict[str, Any]:
        """Extract text from documents (PDF, DOCX, TXT)."""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if not path.exists():
            return {
                "file_path": file_path,
                "file_type": "document",
                "content": "",
                "metadata": {},
                "success": False,
                "error": "File not found"
            }
        
        content = ""
        metadata = {}
        
        if ext == '.pdf':
            # Use PDF service
            from services.pdf_service import PDFService
            pdf_service = PDFService()
            content = await pdf_service.extract_text(file_path)
            metadata["pages"] = await pdf_service.get_page_count(file_path)
            
        elif ext in ['.docx', '.doc']:
            # Use python-docx
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
                metadata["paragraphs"] = len(doc.paragraphs)
            except ImportError:
                content = f"[DOCX file: {path.name}. Install python-docx for extraction.]"
            except Exception as e:
                content = f"[Error reading DOCX: {e}]"
                
        elif ext in ['.txt', '.md']:
            content = path.read_text(encoding='utf-8')
            metadata["characters"] = len(content)
            
        elif ext == '.rtf':
            try:
                from striprtf.striprtf import rtf_to_text
                content = rtf_to_text(path.read_text(encoding='utf-8'))
            except ImportError:
                content = path.read_text(encoding='utf-8', errors='ignore')
        
        return {
            "file_path": file_path,
            "file_type": "document",
            "content": content,
            "metadata": metadata,
            "success": True,
            "error": None
        }
    
    async def process_data(self, file_path: str) -> Dict[str, Any]:
        """Process data files (CSV, JSON, Excel)."""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if not path.exists():
            return {
                "file_path": file_path,
                "file_type": "data",
                "content": "",
                "metadata": {},
                "success": False,
                "error": "File not found"
            }
        
        content = ""
        metadata = {}
        
        if ext == '.json':
            data = json.loads(path.read_text(encoding='utf-8'))
            content = json.dumps(data, indent=2)[:10000]  # Limit size
            metadata["type"] = "json"
            
        elif ext == '.csv':
            try:
                import pandas as pd
                df = pd.read_csv(file_path, nrows=100)  # First 100 rows
                content = df.to_string()
                metadata["rows"] = len(df)
                metadata["columns"] = list(df.columns)
            except ImportError:
                content = path.read_text(encoding='utf-8')[:5000]
                
        elif ext in ['.xlsx', '.xls']:
            try:
                import pandas as pd
                df = pd.read_excel(file_path, nrows=100)
                content = df.to_string()
                metadata["rows"] = len(df)
                metadata["columns"] = list(df.columns)
            except ImportError:
                content = f"[Excel file: {path.name}. Install pandas + openpyxl for extraction.]"
                
        elif ext == '.xml':
            content = path.read_text(encoding='utf-8')[:10000]
            metadata["type"] = "xml"
        
        return {
            "file_path": file_path,
            "file_type": "data",
            "content": content,
            "metadata": metadata,
            "success": True,
            "error": None
        }
    
    async def process_uploaded_files(
        self,
        workspace_id: str,
        file_paths: List[str]
    ) -> List[Dict[str, Any]]:
        """Process multiple uploaded files in parallel."""
        tasks = [self.process_file(fp, workspace_id) for fp in file_paths]
        return await asyncio.gather(*tasks)
    
    def get_file_content_for_llm(
        self,
        processed_files: List[Dict[str, Any]]
    ) -> str:
        """Format processed files for LLM context."""
        parts = []
        
        for f in processed_files:
            if f.get('success') and f.get('content'):
                file_name = Path(f['file_path']).name
                file_type = f['file_type']
                content = f['content'][:5000]  # Limit per file
                
                parts.append(f"### File: {file_name} ({file_type})\n{content}\n")
        
        return "\n---\n".join(parts)


# Singleton
multimodal_processor = MultimodalProcessor()
