"""
Document Export Service

Exports literature synthesis and other documents to various formats:
- DOCX with clickable citation links (internal bookmarks)
- PDF
- HTML

Handles proper citation hyperlinking for Word documents.
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Install with: pip install python-docx")


class DocumentExporter:
    """
    Export documents to DOCX/PDF with proper citation hyperlinking.
    """
    
    def __init__(self):
        self.citation_pattern = re.compile(
            r'\(([A-Za-z\-]+(?:\s+(?:&|and)\s+[A-Za-z\-]+)?(?:\s+et\s+al\.)?),?\s*(\d{4})\)'
        )
    
    def export_to_docx(
        self,
        content: str,
        output_path: str,
        sources: List[Dict[str, Any]],
        title: str = "Literature Synthesis",
        include_toc: bool = True
    ) -> str:
        """
        Export markdown content to DOCX with clickable citation links.
        
        Citations like (Smith, 2020) become hyperlinks that jump to
        the corresponding reference in the References section.
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        doc = Document()
        
        # Build citation key mapping
        citation_map = self._build_citation_map(sources)
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").italic = True
        meta.add_run(f"Based on {len(sources)} sources").italic = True
        
        doc.add_paragraph()  # Spacer
        
        # Process content by sections
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Handle headings
            if line.startswith('# ') and 'Literature Synthesis' not in line:
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                section_title = line[3:]
                if section_title == 'References':
                    # Add page break before references
                    doc.add_page_break()
                doc.add_heading(section_title, 2)
                current_section = section_title
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('---'):
                # Horizontal rule - skip
                continue
            elif line.startswith('*Generated on') or line.startswith('*Based on'):
                # Skip metadata (already added)
                continue
            elif line.startswith('<a id="ref-'):
                # This is a reference anchor - extract key and add bookmark
                import re
                match = re.search(r'id="(ref-[^"]+)"', line)
                if match:
                    bookmark_name = match.group(1)
                    # Next line will be the reference
                    continue
            elif current_section == 'References' and line.startswith('**['):
                # Reference entry - add with bookmark
                self._add_reference_with_bookmark(doc, line, citation_map)
            else:
                # Regular paragraph - process citations
                self._add_paragraph_with_citations(doc, line, citation_map)
        
        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _build_citation_map(self, sources: List[Dict]) -> Dict[str, Dict]:
        """Build mapping from author-year patterns to source info."""
        mapping = {}
        
        for source in sources:
            citation_key = source.get('citation_key', '')
            authors = source.get('authors', [])
            year = source.get('year', '')
            
            if not citation_key or not year:
                continue
            
            # Get author last names
            if authors:
                first = authors[0]
                first_name = first.get('name', first) if isinstance(first, dict) else first
                first_last = first_name.split()[-1] if first_name else ''
                
                if len(authors) == 1:
                    pattern = f"{first_last}, {year}"
                    mapping[pattern] = {
                        'key': citation_key,
                        'bookmark': f"ref_{citation_key}",
                        'doi': source.get('doi', ''),
                        'url': source.get('url', '')
                    }
                elif len(authors) == 2:
                    second = authors[1]
                    second_name = second.get('name', second) if isinstance(second, dict) else second
                    second_last = second_name.split()[-1] if second_name else ''
                    
                    for connector in ['&', 'and']:
                        pattern = f"{first_last} {connector} {second_last}, {year}"
                        mapping[pattern] = {
                            'key': citation_key,
                            'bookmark': f"ref_{citation_key}",
                            'doi': source.get('doi', ''),
                            'url': source.get('url', '')
                        }
                else:
                    pattern = f"{first_last} et al., {year}"
                    mapping[pattern] = {
                        'key': citation_key,
                        'bookmark': f"ref_{citation_key}",
                        'doi': source.get('doi', ''),
                        'url': source.get('url', '')
                    }
        
        return mapping
    
    def _add_paragraph_with_citations(
        self, 
        doc: 'Document', 
        text: str, 
        citation_map: Dict
    ):
        """Add paragraph with clickable citation hyperlinks."""
        para = doc.add_paragraph()
        
        # Find all citations in the text
        last_end = 0
        
        for match in self.citation_pattern.finditer(text):
            # Add text before citation
            if match.start() > last_end:
                para.add_run(text[last_end:match.start()])
            
            # Get citation text
            author_part = match.group(1)
            year = match.group(2)
            citation_text = f"({author_part}, {year})"
            
            # Look up in mapping
            lookup_key = f"{author_part}, {year}"
            if lookup_key in citation_map:
                # Add as internal hyperlink
                self._add_internal_hyperlink(
                    para, 
                    citation_text,
                    citation_map[lookup_key]['bookmark']
                )
            else:
                # Just add as regular text
                para.add_run(citation_text)
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            para.add_run(text[last_end:])
    
    def _add_internal_hyperlink(self, paragraph, text: str, bookmark_name: str):
        """Add an internal hyperlink that jumps to a bookmark."""
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        # Create run with text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Blue color and underline for hyperlink style
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        new_run.append(rPr)
        
        # Add text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def _add_reference_with_bookmark(
        self, 
        doc: 'Document', 
        line: str, 
        citation_map: Dict
    ):
        """Add reference entry with a bookmark anchor."""
        # Extract citation key
        key_match = re.search(r'\*\*\[([^\]]+)\]\*\*', line)
        if not key_match:
            doc.add_paragraph(line)
            return
        
        citation_key = key_match.group(1)
        bookmark_name = f"ref_{citation_key}"
        
        # Create paragraph
        para = doc.add_paragraph()
        
        # Add bookmark at start
        self._add_bookmark(para, bookmark_name)
        
        # Parse reference parts
        # Format: **[key]** Authors (Year). *Title*. [URL]
        line = line.replace('**', '').replace('*', '')
        
        # Add the reference text
        run = para.add_run(line)
        
        # Add hanging indent style
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
    
    def _add_bookmark(self, paragraph, bookmark_name: str):
        """Add a bookmark anchor to a paragraph."""
        # Create bookmark start
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        # Create bookmark end
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')
        
        # Add to paragraph
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)
    
    def export_synthesis_to_docx(
        self,
        workspace_id: str,
        synthesis_content: str,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Export a synthesis to DOCX with proper citation links.
        """
        from services.sources_service import sources_service
        from services.workspace_service import WORKSPACES_DIR
        
        # Get sources for citation mapping
        sources = sources_service.list_sources(workspace_id)
        
        # Generate output path
        if not output_filename:
            output_filename = f"synthesis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        output_dir = WORKSPACES_DIR / workspace_id / "outputs"
        output_path = output_dir / output_filename
        
        # Export
        return self.export_to_docx(
            content=synthesis_content,
            output_path=str(output_path),
            sources=sources,
            title="Literature Synthesis"
        )


# Singleton
document_exporter = DocumentExporter()
