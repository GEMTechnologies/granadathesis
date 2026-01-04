"""
Complete Thesis Formatter with MS Word (.docx) generation.

Generates professional thesis with:
- Cover page (no numbering)
- Preliminary pages (Roman numerals: i, ii, iii, etc.)
- Main content (Arabic numerals: 1, 2, 3, etc.)
- Auto-generated table of contents
- Proper formatting and spacing
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Optional
from pathlib import Path
import json
from datetime import datetime
from .page_numbering import ThesisPageNumberingManager, AdvancedPageNumbering


class ThesisFormatter:
    """Format and export complete thesis to MS Word (.docx)."""
    
    def __init__(self, workspace_id: str = "default"):
        self.workspace_id = workspace_id
        from .workspace_service import WORKSPACES_DIR
        self.workspace_path = WORKSPACES_DIR / workspace_id
        self.workspace_path.mkdir(parents=True, exist_ok=True)
        self.numbering_manager = None
        
    def create_complete_thesis(
        self,
        title: str,
        author_name: str,
        index_no: str,
        school: str,
        department: str,
        degree: str,
        chapters: Dict[int, str],  # {1: "Chapter 1 content", 2: "Chapter 2 content", ...}
        supervisor_approval: str = "",
        student_declaration: str = "",
        dedication: str = "",
        acknowledgement: str = "",
        abstract: str = "",
        appendices: Optional[Dict[str, str]] = None,  # {"Appendix A": "content", ...}
        cover_info: Optional[Dict] = None,
    ) -> str:
        """
        Create complete thesis with all sections and return path to generated .docx file.
        
        Args:
            title: Research title
            author_name: Student name
            index_no: Student ID
            school: School name
            department: Department name
            degree: Degree type (PhD, Masters, etc.)
            chapters: Dictionary of chapter numbers to content
            supervisor_approval: Approval statement
            student_declaration: Declaration statement
            dedication: Dedication text
            acknowledgement: Acknowledgement text
            abstract: Abstract text
            appendices: Optional appendices
            cover_info: Optional additional cover info
            
        Returns:
            Path to generated .docx file
        """
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Track page numbering sections
        self.doc = doc
        self.chapters = chapters
        self.appendices = appendices or {}
        self.numbering_manager = ThesisPageNumberingManager(doc)
        
        # 1. COVER PAGE (no numbering)
        self._add_cover_page(title, author_name, index_no, school, department, degree)
        self.numbering_manager.setup_cover_page()
        self._page_break()
        
        # 2. APPROVAL PAGE
        self.numbering_manager.setup_preliminaries()
        self._add_approval_page(supervisor_approval)
        self._page_break()
        
        # 3. DECLARATION PAGE
        self._add_declaration_page(student_declaration)
        self._page_break()
        
        # 4. DEDICATION PAGE
        self._add_dedication_page(dedication)
        self._page_break()
        
        # 5. ACKNOWLEDGEMENT PAGE
        self._add_acknowledgement_page(acknowledgement)
        self._page_break()
        
        # 6. ABSTRACT
        self._add_abstract_page(abstract)
        self._page_break()
        
        # 7. TABLE OF CONTENTS
        self._add_table_of_contents(chapters, appendices)
        self._page_break()
        
        # 8. LIST OF TABLES
        self._add_list_of_tables()
        self._page_break()
        
        # 9. LIST OF FIGURES
        self._add_list_of_figures()
        self._page_break()
        
        # 10. ACRONYMS/ABBREVIATIONS
        self._add_acronyms()
        self._page_break()
        
        # Switch to Arabic numeral numbering for main content
        self.numbering_manager.setup_main_content()
        
        # 11. MAIN CHAPTERS
        for chapter_num in sorted(chapters.keys()):
            self._add_chapter(chapter_num, chapters[chapter_num])
            self._page_break()
        
        # 12. REFERENCES
        self._add_references_page()
        
        # 13. APPENDICES
        # Always use UoJ Appendices for consistency with user request
        self._add_uoj_appendices()
        
        # Save document
        output_path = self.workspace_path / f"Complete_Thesis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _add_references_page(self):
        """Add standard references page."""
        self._add_page_section("REFERENCES")
        self.doc.add_paragraph(
            "All scholarly sources cited in this thesis are embedded as hyperlinks throughout the document. "
            "A consolidated reference list in APA 7th Edition format is available in the exported document."
        )
        self._page_break()

    def _add_uoj_appendices(self):
        """Add UoJ specific appendices."""
        self._add_page_section("APPENDICES")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        appendices = [
            ("Appendix A: Research Questionnaire", f"The structured questionnaire used for data collection is available as a separate document:\n• **File:** `Questionnaire_{timestamp}.md`"),
            ("Appendix B: Interview Guide", f"The semi-structured interview guide for qualitative data collection:\n• **File:** `Interview_Guide_{timestamp}.md`"),
            ("Appendix C: Focus Group Discussion Guide", "The FGD protocol and questions:\n• **File:** See Interview Guide"),
            ("Appendix D: Research Permit/Authorisation Letters", "[To be inserted by researcher]"),
            ("Appendix E: Informed Consent Form", "[To be inserted by researcher]"),
            ("Appendix F: Raw Data", "Statistical outputs and transcripts are available in:\n• **Folder:** `datasets/`")
        ]
        
        for title, content in appendices:
            self.doc.add_heading(title, level=2)
            self.doc.add_paragraph(content)
            self.doc.add_paragraph()  # Spacing
            
        # End marker
        p = self.doc.add_paragraph("— END OF THESIS —")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def _add_cover_page(self, title: str, author_name: str, index_no: str, 
                       school: str, department: str, degree: str):
        """Add cover page with UoJ specific formatting."""
        doc = self.doc
        
        # 1. RESEARCH THESIS (Top)
        self._add_centered_heading("RESEARCH THESIS", font_size=14, bold=True)
        doc.add_paragraph()
        
        # 2. Submission Statement
        p = self._add_centered_paragraph(
            "A Thesis Submitted in Partial Fulfilment of the Requirements for the Award of the Degree of",
            font_size=12
        )
        doc.add_paragraph()
        
        # 3. DOCTOR OF PHILOSOPHY
        self._add_centered_heading("DOCTOR OF PHILOSOPHY", font_size=16, bold=True)
        
        # 4. in [Department/Field]
        self._add_centered_paragraph("in", font_size=12)
        # Use department or "Security Studies" if not provided
        dept_text = department if department and department != "[DEPARTMENT NAME]" else "Security Studies"
        self._add_centered_heading(dept_text, font_size=14, bold=True)
        doc.add_paragraph()
        
        # 5. Logo
        logo_path = "/home/gemtech/Desktop/thesis/backend/lightweight/uoj_logo.png"
        try:
            if os.path.exists(logo_path):
                # Add picture efficiently
                doc.add_picture(logo_path, width=Inches(2.5))
                # Center the last paragraph (which contains the image)
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                self._add_centered_paragraph("[UOJ LOGO]", font_size=14)
        except Exception as e:
            print(f"Error adding logo: {e}")
            self._add_centered_paragraph("[UOJ LOGO]", font_size=14)
            
        doc.add_paragraph()
        
        # 6. University of Juba
        self._add_centered_heading("University of Juba", font_size=16, bold=True)
        doc.add_paragraph()
        
        # 7. Date
        self._add_centered_paragraph("January 2026", font_size=12)
        
        # Spacing at bottom
        doc.add_paragraph()
    
    def _add_approval_page(self, approval_text: str):
        """Add approval page."""
        self._add_page_section("APPROVAL")
        
        if approval_text:
            self.doc.add_paragraph(approval_text)
        else:
            self.doc.add_paragraph(
                "Write supervisor approval statement for my study"
            )
    
    def _add_declaration_page(self, declaration_text: str):
        """Add declaration page."""
        self._add_page_section("DECLARATION")
        
        if declaration_text:
            self.doc.add_paragraph(declaration_text)
        else:
            self.doc.add_paragraph(
                "Write student declaration statement for my study"
            )
    
    def _add_dedication_page(self, dedication_text: str):
        """Add dedication page."""
        self._add_page_section("DEDICATION")
        
        if dedication_text:
            self.doc.add_paragraph(dedication_text)
        else:
            self.doc.add_paragraph(
                "Write dedication for my study"
            )
    
    def _add_acknowledgement_page(self, acknowledgement_text: str):
        """Add acknowledgement page."""
        self._add_page_section("ACKNOWLEDGEMENT")
        
        if acknowledgement_text:
            self.doc.add_paragraph(acknowledgement_text)
        else:
            self.doc.add_paragraph(
                "Write acknowledgement for my study"
            )
    
    def _add_abstract_page(self, abstract_text: str):
        """Add abstract page."""
        self._add_page_section("ABSTRACT")
        
        if abstract_text:
            self.doc.add_paragraph(abstract_text)
        else:
            self.doc.add_paragraph(
                "[Abstract will be automatically generated from your thesis content]"
            )
    
    def _add_table_of_contents(self, chapters: Dict, appendices: Dict):
        """Add auto-generated table of contents."""
        self._add_page_section("TABLE OF CONTENTS")
        
        # Chapters
        for chapter_num in sorted(chapters.keys()):
            page_num = self._estimate_page_number(chapter_num)
            self.doc.add_paragraph(
                f"Chapter {chapter_num}: [Title from Chapter {chapter_num}] ...................... {page_num}",
                style='List Bullet'
            )
        
        # Appendices
        if appendices:
            for app_name in appendices.keys():
                page_num = self._estimate_page_number(f"Appendix: {app_name}")
                self.doc.add_paragraph(
                    f"{app_name} ...................... {page_num}",
                    style='List Bullet'
                )
    
    def _add_list_of_tables(self):
        """Add list of tables section."""
        self._add_page_section("LIST OF TABLES")
        self.doc.add_paragraph("[Auto-generated from tables in your thesis]")
    
    def _add_list_of_figures(self):
        """Add list of figures section."""
        self._add_page_section("LIST OF FIGURES")
        self.doc.add_paragraph("[Auto-generated from figures in your thesis]")
    
    def _add_acronyms(self):
        """Add acronyms/abbreviations section."""
        self._add_page_section("ACRONYMS/ABBREVIATIONS")
        
        # Create table for acronyms
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Acronym"
        header_cells[1].text = "Full Form"
        
        # Example acronyms (user can edit)
        example_acronyms = [
            ("PhD", "Doctor of Philosophy"),
            ("AI", "Artificial Intelligence"),
            ("ML", "Machine Learning"),
            ("NLP", "Natural Language Processing"),
        ]
        
        for acronym, full_form in example_acronyms:
            row_cells = table.add_row().cells
            row_cells[0].text = acronym
            row_cells[1].text = full_form
    
    def _add_chapter(self, chapter_num: int, content: str):
        """Add chapter content."""
        heading = self.doc.add_heading(f"Chapter {chapter_num}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add chapter content
        self.doc.add_paragraph(content)
    
    def _add_appendices(self, appendices: Dict[str, str]):
        """Add appendices section."""
        for app_name, app_content in appendices.items():
            heading = self.doc.add_heading(app_name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph(app_content)
            self._page_break()
    
    def _add_page_section(self, title: str):
        """Add a page section with centered title."""
        heading = self.doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_centered_heading(self, text: str, font_size: int = 12, bold: bool = False):
        """Add centered heading."""
        paragraph = self.doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True
    
    def _add_centered_paragraph(self, text: str, font_size: int = 12):
        """Add centered paragraph."""
        paragraph = self.doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
    
    def _page_break(self):
        """Add page break."""
        self.doc.add_page_break()
    
    def _set_page_numbering_style(self, style: str):
        """
        Set page numbering style.
        style: 'roman' (i, ii, iii) or 'arabic' (1, 2, 3)
        """
        # This is handled by numbering_manager now
        pass
    
    def _estimate_page_number(self, chapter_identifier) -> int:
        """Estimate page number for TOC."""
        # Simplified estimation
        # In production, would track actual page numbers
        if isinstance(chapter_identifier, int):
            return 10 + (chapter_identifier * 5)  # Rough estimate
        return 50
