"""
Chapter Generator - Generates structured thesis chapters with citations

Creates properly formatted chapters with:
- Centered headings (CHAPTER ONE, INTRODUCTION)
- Numbered sections (1.1, 1.4, 1.5)
- Heavily cited content
- Injected objectives from database
- Reference list
"""

from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from app.services.cited_content_generator import cited_content_generator


class ChapterGenerator:
    """Generates structured thesis chapters with proper formatting."""
    
    def __init__(self, citation_style: str = "APA"):
        self.citation_style = citation_style
        self.content_generator = cited_content_generator
    
    async def generate_chapter_one(
        self,
        topic: str,
        case_study: str,
        objectives: Optional[List[Dict[str, Any]]] = None,
        research_questions: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Generate Chapter 1: Introduction
        
        Args:
            topic: Research topic
            case_study: Case study/context
            objectives: List of objectives from MAKER voting
            research_questions: List of research questions
            
        Returns:
            Dict with content, DOCX document, and metadata
        """
        print("\n" + "="*80)
        print("GENERATING CHAPTER ONE: INTRODUCTION")
        print("="*80 + "\n")
        
        # Initialize document
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add chapter title (centered)
        self._add_centered_heading(doc, "CHAPTER ONE", level=0)
        self._add_centered_heading(doc, "INTRODUCTION", level=0)
        doc.add_paragraph()  # Spacing
        
        # Section 1.1: Setting the Scene (with error handling)
        print("📝 Generating Section 1.1: Setting the Scene...")
        try:
            setting_scene = await self.content_generator.generate_cited_section(
                section_title="Setting the Scene",
                topic=f"{topic} in {case_study}",
                word_count=500,
                target_density=0.75
            )
        except Exception as e:
            print(f"   ⚠️  Error generating cited content: {str(e)[:100]}")
            print("   📝 Falling back to minimal content generation...")
            
            # Fallback: Create minimal section without citations
            setting_scene = {
                'content': (
                    f"This research focuses on {topic} in the context of {case_study}. "
                    f"The study aims to investigate key aspects and contribute to the existing body of knowledge. "
                    f"Further details and citations will be added upon successful API connectivity."
                ),
                'references': [],
                'metrics': {
                    'word_count': 40,
                    'sentence_count': 3,
                    'citation_count': 0,
                    'citation_density': 0.0,
                    'unique_papers': 0
                }
            }
            print("   ✓ Fallback content generated")
        
        self._add_section(
            doc,
            section_number="1.1",
            section_title="Setting the Scene",
            content=setting_scene['content']
        )
        
        # Section 1.4: Objectives
        print("\n📋 Adding Section 1.4: Objectives...")
        self._add_objectives_section(doc, objectives or [])
        
        # Section 1.5: Research Questions
        print("\n❓ Adding Section 1.5: Research Questions...")
        self._add_research_questions_section(doc, research_questions or [])
        
        # References
        print("\n📚 Adding References...")
        all_references = setting_scene.get('references', [])
        self._add_references_section(doc, all_references)
        
        # Compile metadata
        metadata = {
            "chapter_number": 1,
            "chapter_title": "INTRODUCTION",
            "sections": [
                {
                    "number": "1.1",
                    "title": "Setting the Scene",
                    "word_count": setting_scene['metrics']['word_count'],
                    "citation_count": setting_scene['metrics']['citation_count']
                },
                {"number": "1.4", "title": "Objectives"},
                {"number": "1.5", "title": "Research Questions/Hypothesis"}
            ],
            "total_references": len(all_references),
            "metrics": setting_scene['metrics']
        }
        
        print("\n✅ Chapter One Complete!")
        print(f"   Sections: {len(metadata['sections'])}")
        print(f"   References: {metadata['total_references']}")
        print(f"   Citation density: {setting_scene['metrics']['citation_density']:.1%}\n")
        
        return {
            "document": doc,
            "metadata": metadata,
            "content": setting_scene['content'],
            "references": all_references
        }
    
    def _setup_document_styles(self, doc: Document):
        """Setup document-wide styles."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set paragraph formatting
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.space_after = Pt(6)
    
    def _add_centered_heading(
        self,
        doc: Document,
        text: str,
        level: int = 0
    ):
        """Add centered heading."""
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format heading
        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
    
    def _add_section(
        self,
        doc: Document,
        section_number: str,
        section_title: str,
        content: str
    ):
        """Add a numbered section."""
        # Section heading
        heading = doc.add_heading(f"{section_number} {section_title}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Format heading
        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        
        # Section content
        paragraph = doc.add_paragraph(content)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Format content
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    def _add_objectives_section(
        self,
        doc: Document,
        objectives: List[Dict[str, Any]]
    ):
        """Add objectives section with subsections."""
        # Main heading
        heading = doc.add_heading("1.4 Objectives", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        
        # 1.4.1 General/Broad Objectives
        subheading = doc.add_heading("1.4.1 General/Broad Objectives", level=3)
        subheading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = subheading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        
        # Add general objectives
        general_objectives = [obj for obj in objectives if obj.get('type') == 'general']
        if general_objectives:
            for obj in general_objectives:
                p = doc.add_paragraph(obj.get('text', ''), style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p = doc.add_paragraph("To be determined based on research scope.")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 1.4.2 Specific Objectives
        subheading = doc.add_heading("1.4.2 Specific Objectives", level=3)
        subheading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = subheading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        
        # Add specific objectives
        specific_objectives = [obj for obj in objectives if obj.get('type') == 'specific']
        if specific_objectives:
            for obj in specific_objectives:
                p = doc.add_paragraph(obj.get('text', ''), style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p = doc.add_paragraph("To be determined based on research scope.")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_research_questions_section(
        self,
        doc: Document,
        research_questions: List[str]
    ):
        """Add research questions section."""
        # Section heading
        heading = doc.add_heading("1.5 Research Questions/Hypothesis", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        
        # Add research questions
        if research_questions:
            for i, question in enumerate(research_questions, 1):
                p = doc.add_paragraph(f"RQ{i}: {question}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        else:
            p = doc.add_paragraph("Research questions will be formulated based on the literature review and identified research gaps.")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_references_section(
        self,
        doc: Document,
        references: List[str]
    ):
        """Add references section with hanging indent."""
        # Add page break
        doc.add_page_break()
        
        # References heading (centered)
        heading = doc.add_heading("References", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        
        # Add each reference with hanging indent
        for reference in references:
            p = doc.add_paragraph(reference)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Hanging indent
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            
            # Format
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def save_document(self, doc: Document, filename: str):
        """Save document to file."""
        doc.save(filename)
        print(f"✅ Document saved: {filename}")


# Singleton instance
chapter_generator = ChapterGenerator(citation_style="APA")
